"""
Flask API for But-For Damages Analyzer
"""
import hashlib
import io
import json
import os
import tempfile
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from sqlalchemy.exc import IntegrityError
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import openpyxl
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import numpy as np

from models import db, init_db, Evaluee, Case, Calculation
from config import config


# Range validation rules to enforce accuracy and block implausible scenarios
RANGE_VALIDATION_RULES = [
    {
        'label': 'Growth rate',
        'path': ('butFor', 'growth'),
        'min': -0.05,
        'max': 0.15,
        'description': 'Expected to fall between -5% (contraction) and 15% (aggressive growth).'
    },
    {
        'label': 'Discount rate',
        'path': ('discount', 'rate'),
        'min': -0.05,
        'max': 0.2,
        'description': 'Net discount rates should typically remain within -5% and 20%.'
    },
    {
        'label': 'Net discount rate',
        'path': ('discount', 'ndr'),
        'min': -0.05,
        'max': 0.2,
        'description': 'Net discount rates should typically remain within -5% and 20%.'
    },
    {
        'label': 'Unemployment factor',
        'path': ('aef', 'ufEff'),
        'min': 0,
        'max': 0.35,
        'description': 'Unemployment factors above 35% or below 0 are flagged as implausible.'
    },
    {
        'label': 'Tax load',
        'path': ('aef', 'tlEff'),
        'min': 0,
        'max': 0.65,
        'description': 'Combined effective tax loads rarely exceed 65% of wages.'
    },
    {
        'label': 'Fringe/benefit percentage',
        'path': ('aef', 'fringePct'),
        'min': 0,
        'max': 0.75,
        'description': 'Fringe loads above 75% of wages should be reviewed before use.'
    },
]


def _get_nested_value(data, path):
    """Safely walk a dict using a tuple path."""

    cursor = data if isinstance(data, dict) else {}
    for key in path:
        if not isinstance(cursor, dict):
            return None
        cursor = cursor.get(key)
    return cursor


def _validate_assumption_ranges(assumptions):
    """Validate assumption values against configured ranges.

    Returns a list of human-readable violations. An empty list indicates all
    monitored fields are within range or unavailable.
    """

    violations = []
    for rule in RANGE_VALIDATION_RULES:
        raw_value = _get_nested_value(assumptions, rule['path'])
        if raw_value is None:
            continue

        try:
            value = float(raw_value)
        except (TypeError, ValueError):
            violations.append(
                f"{rule['label']}: expected a numeric value but received {raw_value!r}."
            )
            continue

        if not (rule['min'] <= value <= rule['max']):
            violations.append(
                f"{rule['label']} {value:.4g} is outside the expected range "
                f"[{rule['min']}, {rule['max']}]. {rule['description']}"
            )

    return violations


def _compute_assumption_fingerprint(assumptions):
    """Generate a stable SHA-256 fingerprint for the provided assumptions."""

    try:
        normalized = json.dumps(assumptions or {}, sort_keys=True, default=str)
    except TypeError:
        normalized = json.dumps(str(assumptions or {}), sort_keys=True)
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()


def _build_provenance_metadata(assumptions):
    """Create provenance metadata including sources, fingerprint, and timestamp."""

    meta = assumptions.get('meta', {}) if isinstance(assumptions, dict) else {}
    life_table = assumptions.get('lifeTable', {}) if isinstance(assumptions, dict) else {}

    sources = []
    if life_table.get('source'):
        population = life_table.get('population', 'combined').title()
        sources.append(f"Life table: {life_table['source']} [{population}]")
    if meta.get('wageSourceNotes'):
        sources.append(f"Wage/growth documentation: {meta['wageSourceNotes']}")
    if meta.get('benefitSourceNotes'):
        sources.append(f"Fringe/benefit documentation: {meta['benefitSourceNotes']}")

    return {
        'fingerprint': _compute_assumption_fingerprint(assumptions),
        'generated_at': datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'),
        'sources': sources,
    }


def _parse_date(date_value, field_name):
    """Convert an ISO date string to a ``datetime.date`` while validating input."""

    if not date_value:
        return None

    try:
        return datetime.fromisoformat(date_value).date()
    except ValueError:
        raise ValueError(f"Invalid date format for '{field_name}'. Use ISO format YYYY-MM-DD.")


def _handle_integrity_error(error: IntegrityError):
    db.session.rollback()
    return jsonify({'success': False, 'error': 'Data integrity error', 'details': str(error.orig)}), 400


def create_app(config_name='development'):
    """Create and configure Flask app."""
    app = Flask(__name__)

    # Load config
    app.config.from_object(config[config_name])

    # Ensure data directory exists
    data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data')
    os.makedirs(data_dir, exist_ok=True)

    # Initialize extensions
    CORS(app, origins=app.config['CORS_ORIGINS'])
    init_db(app)

    # ==================== EVALUEE ENDPOINTS ====================

    @app.route('/api/evaluees', methods=['GET'])
    def get_evaluees():
        """Get all evaluees"""
        evaluees = Evaluee.query.order_by(Evaluee.profile_name).all()
        return jsonify({
            'success': True,
            'evaluees': [e.to_dict() for e in evaluees]
        })

    @app.route('/api/evaluees/<int:evaluee_id>', methods=['GET'])
    def get_evaluee(evaluee_id):
        """Get specific evaluee with their cases"""
        evaluee = Evaluee.query.get_or_404(evaluee_id)
        data = evaluee.to_dict()
        data['cases'] = [c.to_dict(include_assumptions=False) for c in evaluee.cases]
        return jsonify({
            'success': True,
            'evaluee': data
        })

    @app.route('/api/evaluees', methods=['POST'])
    def create_evaluee():
        """Create new evaluee"""
        data = request.json
        profile_name = data.get('profile_name', '').strip()

        if not profile_name:
            return jsonify({'success': False, 'error': 'Profile name is required'}), 400

        # Check for duplicates
        existing = Evaluee.query.filter_by(profile_name=profile_name).first()
        if existing:
            return jsonify({'success': False, 'error': 'Profile name already exists'}), 400

        evaluee = Evaluee(profile_name=profile_name)
        db.session.add(evaluee)
        try:
            db.session.commit()
        except IntegrityError as exc:  # Defensive against race-conditions
            return _handle_integrity_error(exc)

        return jsonify({
            'success': True,
            'evaluee': evaluee.to_dict()
        }), 201

    @app.route('/api/evaluees/<int:evaluee_id>', methods=['PUT'])
    def update_evaluee(evaluee_id):
        """Update evaluee"""
        evaluee = Evaluee.query.get_or_404(evaluee_id)
        data = request.json
        profile_name = data.get('profile_name', '').strip()

        if profile_name:
            # Check for duplicates (excluding self)
            existing = Evaluee.query.filter(
                Evaluee.profile_name == profile_name,
                Evaluee.id != evaluee_id
            ).first()
            if existing:
                return jsonify({'success': False, 'error': 'Profile name already exists'}), 400

            evaluee.profile_name = profile_name

        evaluee.updated_at = datetime.utcnow()
        try:
            db.session.commit()
        except IntegrityError as exc:
            return _handle_integrity_error(exc)

        return jsonify({
            'success': True,
            'evaluee': evaluee.to_dict()
        })

    @app.route('/api/evaluees/<int:evaluee_id>', methods=['DELETE'])
    def delete_evaluee(evaluee_id):
        """Delete evaluee and all associated cases"""
        evaluee = Evaluee.query.get_or_404(evaluee_id)
        db.session.delete(evaluee)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Evaluee "{evaluee.profile_name}" deleted'
        })

    # ==================== CASE ENDPOINTS ====================

    @app.route('/api/evaluees/<int:evaluee_id>/cases', methods=['GET'])
    def get_cases(evaluee_id):
        """Get all cases for an evaluee"""
        evaluee = Evaluee.query.get_or_404(evaluee_id)
        cases = Case.query.filter_by(evaluee_id=evaluee_id).order_by(Case.updated_at.desc()).all()

        return jsonify({
            'success': True,
            'evaluee': evaluee.to_dict(),
            'cases': [c.to_dict(include_assumptions=True) for c in cases]
        })

    @app.route('/api/cases/<int:case_id>', methods=['GET'])
    def get_case(case_id):
        """Get specific case with full details"""
        case = Case.query.get_or_404(case_id)
        include_history = request.args.get('include_history', 'false').lower() == 'true'

        return jsonify({
            'success': True,
            'case': case.to_dict(include_assumptions=True, include_calculations=include_history)
        })

    @app.route('/api/evaluees/<int:evaluee_id>/cases', methods=['POST'])
    def create_case(evaluee_id):
        """Create new case for evaluee"""
        evaluee = Evaluee.query.get_or_404(evaluee_id)
        data = request.json

        try:
            case = Case(
                evaluee_id=evaluee_id,
                case_name=data.get('case_name', 'Untitled Case').strip() or 'Untitled Case',
                case_type=data.get('case_type', 'pi'),
                date_of_birth=_parse_date(data.get('date_of_birth'), 'date_of_birth'),
                incident_date=_parse_date(data.get('incident_date'), 'incident_date'),
                valuation_date=_parse_date(data.get('valuation_date'), 'valuation_date'),
                wle_years=data.get('wle_years'),
                yfs_years=data.get('yfs_years'),
                le_years=data.get('le_years'),
                assumptions=data.get('assumptions') or {},
            )
        except ValueError as exc:
            return jsonify({'success': False, 'error': str(exc)}), 400

        db.session.add(case)
        try:
            db.session.commit()
        except IntegrityError as exc:
            return _handle_integrity_error(exc)

        return jsonify({
            'success': True,
            'case': case.to_dict()
        }), 201

    @app.route('/api/cases/<int:case_id>', methods=['PUT'])
    def update_case(case_id):
        """Update case"""
        case = Case.query.get_or_404(case_id)
        data = request.json

        try:
            # Update fields if provided
            if 'case_name' in data:
                case.case_name = data['case_name'].strip() or case.case_name
            if 'case_type' in data:
                case.case_type = data['case_type']
            if 'date_of_birth' in data:
                case.date_of_birth = _parse_date(data.get('date_of_birth'), 'date_of_birth')
            if 'incident_date' in data:
                case.incident_date = _parse_date(data.get('incident_date'), 'incident_date')
            if 'valuation_date' in data:
                case.valuation_date = _parse_date(data.get('valuation_date'), 'valuation_date')
            if 'wle_years' in data:
                case.wle_years = data['wle_years']
            if 'yfs_years' in data:
                case.yfs_years = data['yfs_years']
            if 'le_years' in data:
                case.le_years = data['le_years']
            if 'assumptions' in data and isinstance(data.get('assumptions'), dict):
                case.assumptions = data['assumptions']
            if 'latest_calculation' in data and isinstance(data.get('latest_calculation'), dict):
                case.latest_calculation = data['latest_calculation']
        except ValueError as exc:
            return jsonify({'success': False, 'error': str(exc)}), 400

        case.updated_at = datetime.utcnow()
        try:
            db.session.commit()
        except IntegrityError as exc:
            return _handle_integrity_error(exc)

        return jsonify({
            'success': True,
            'case': case.to_dict()
        })

    @app.route('/api/cases/<int:case_id>', methods=['DELETE'])
    def delete_case(case_id):
        """Delete case"""
        case = Case.query.get_or_404(case_id)
        db.session.delete(case)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Case "{case.case_name}" deleted'
        })

    # ==================== CALCULATION ENDPOINTS ====================

    @app.route('/api/cases/<int:case_id>/calculations', methods=['GET'])
    def get_calculations(case_id):
        """Get calculation history for a case"""
        case = Case.query.get_or_404(case_id)
        calculations = Calculation.query.filter_by(case_id=case_id).order_by(Calculation.calculated_at.desc()).all()

        return jsonify({
            'success': True,
            'case_id': case_id,
            'calculations': [c.to_dict(include_full_results=False) for c in calculations]
        })

    @app.route('/api/calculations/<int:calc_id>', methods=['GET'])
    def get_calculation(calc_id):
        """Get specific calculation with full details"""
        calculation = Calculation.query.get_or_404(calc_id)

        return jsonify({
            'success': True,
            'calculation': calculation.to_dict(include_full_results=True)
        })

    @app.route('/api/cases/<int:case_id>/calculations', methods=['POST'])
    def save_calculation(case_id):
        """Save a calculation result"""
        case = Case.query.get_or_404(case_id)
        data = request.json

        assumptions = data.get('assumptions') or {}
        violations = _validate_assumption_ranges(assumptions)
        if violations:
            return jsonify({
                'success': False,
                'error': 'Assumptions validation failed',
                'violations': violations,
            }), 400

        calculation = Calculation(
            case_id=case_id,
            description=data.get('description'),
            assumptions=assumptions,
            results=data.get('results') or {},
            total_damages_pv=data.get('total_damages_pv'),
            past_damages=data.get('past_damages'),
            future_damages_pv=data.get('future_damages_pv')
        )

        db.session.add(calculation)

        # Update case's latest calculation
        case.latest_calculation = data.get('results') or {}
        case.updated_at = datetime.utcnow()

        try:
            db.session.commit()
        except IntegrityError as exc:
            return _handle_integrity_error(exc)

        return jsonify({
            'success': True,
            'calculation': calculation.to_dict(include_full_results=True)
        }), 201

    @app.route('/api/calculations/<int:calc_id>', methods=['DELETE'])
    def delete_calculation(calc_id):
        """Delete a calculation"""
        calculation = Calculation.query.get_or_404(calc_id)
        db.session.delete(calculation)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Calculation deleted'
        })

    # ==================== UTILITY ENDPOINTS ====================

    @app.route('/api/search', methods=['GET'])
    def search():
        """Search evaluees and cases"""
        query = request.args.get('q', '').strip()
        if not query:
            return jsonify({'success': False, 'error': 'Query parameter required'}), 400

        # Search evaluees
        evaluees = (
            Evaluee.query.filter(Evaluee.profile_name.ilike(f'%{query}%'))
            .order_by(Evaluee.profile_name.asc())
            .limit(25)
            .all()
        )

        # Search cases
        cases = (
            Case.query.filter(Case.case_name.ilike(f'%{query}%'))
            .order_by(Case.updated_at.desc())
            .limit(25)
            .all()
        )

        return jsonify({
            'success': True,
            'query': query,
            'results': {
                'evaluees': [e.to_dict() for e in evaluees],
                'cases': [c.to_dict(include_assumptions=False) for c in cases]
            }
        })

    # ==================== CHART GENERATION HELPERS ====================

    def create_annual_loss_chart(rows, title="Annual Economic Loss Over Time"):
        """Create bar chart showing annual losses"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        losses = [row.get('loss', 0) for row in rows]

        colors = ['#d62728' if loss < 0 else '#2ca02c' for loss in losses]
        bars = ax.bar(years, losses, color=colors, alpha=0.7, edgecolor='black')

        ax.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Annual Loss ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3, axis='y')

        # Format y-axis as currency
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        # Rotate x-axis labels if many years
        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        # Save to bytes
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_earnings_comparison_chart(rows, title="But-For vs Actual Earnings"):
        """Create line chart comparing but-for and actual earnings"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        bf_gross = [row.get('bfGross', 0) for row in rows]
        act_earnings = [row.get('actE', 0) for row in rows]

        ax.plot(years, bf_gross, marker='o', linewidth=2, label='But-For Earnings', color='#2ca02c')
        ax.plot(years, act_earnings, marker='s', linewidth=2, label='Actual Earnings', color='#d62728')

        ax.fill_between(years, bf_gross, act_earnings, alpha=0.2, color='gray')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Earnings ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3)

        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_damages_breakdown_pie(totals, title="Total Damages Breakdown"):
        """Create pie chart showing past vs future damages"""
        if not totals:
            return None

        past_dam = totals.get('pastDam', 0)
        future_pv = totals.get('futurePV', 0)

        if past_dam <= 0 and future_pv <= 0:
            return None

        fig, ax = plt.subplots(figsize=(10, 8))

        labels = ['Past Damages', 'Future Damages (PV)']
        sizes = [past_dam, future_pv]
        colors = ['#ff9999', '#66b3ff']
        explode = (0.05, 0.05)

        wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                                           autopct='%1.1f%%', shadow=True, startangle=90,
                                           textprops={'fontsize': 12, 'fontweight': 'bold'})

        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(14)

        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)

        # Add total in center
        total_pv = totals.get('totalPV', 0)
        ax.text(0, 0, f'Total PV\n${total_pv:,.0f}', ha='center', va='center',
                fontsize=12, fontweight='bold',
                bbox=dict(boxstyle='round', facecolor='white', edgecolor='black', linewidth=2))

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_fringe_benefits_chart(rows, use_ups_fringe, title="Fringe Benefits Analysis"):
        """Create stacked bar chart showing fringe benefits breakdown"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]

        if use_ups_fringe:
            hw = [row.get('bfHW', 0) for row in rows]
            pension = [row.get('bfPension', 0) for row in rows]

            ax.bar(years, hw, label='Health & Welfare', color='#8c564b', alpha=0.8)
            ax.bar(years, pension, bottom=hw, label='Pension', color='#e377c2', alpha=0.8)
        else:
            fringe = [row.get('bfFringe', 0) for row in rows]
            ax.bar(years, fringe, label='Total Fringe', color='#9467bd', alpha=0.8)

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Fringe Benefits ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3, axis='y')

        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_retirement_scenarios_chart(scenarios, title="Retirement Age Scenarios Comparison"):
        """Create grouped bar chart comparing retirement scenarios"""
        if not scenarios or len(scenarios) == 0:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))

        scenario_names = [s.get('name', '') for s in scenarios]
        past_damages = [s.get('totals', {}).get('pastDam', 0) for s in scenarios]
        future_damages = [s.get('totals', {}).get('futurePV', 0) for s in scenarios]

        x = np.arange(len(scenario_names))
        width = 0.35

        bars1 = ax.bar(x - width/2, past_damages, width, label='Past Damages', color='#ff9999', edgecolor='black')
        bars2 = ax.bar(x + width/2, future_damages, width, label='Future Damages (PV)', color='#66b3ff', edgecolor='black')

        ax.set_xlabel('Retirement Scenario', fontsize=12, fontweight='bold')
        ax.set_ylabel('Damages ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(scenario_names, rotation=0)
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3, axis='y')

        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        # Add value labels on bars
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'${height:,.0f}',
                           ha='center', va='bottom', fontsize=8, rotation=0)

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_cumulative_damages_chart(rows, title="Cumulative Damages Over Time"):
        """Create area chart showing cumulative damages"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]

        # Calculate cumulative values
        cumulative_past = []
        cumulative_future = []
        running_past = 0
        running_future = 0

        for row in rows:
            running_past += row.get('pastPart', 0)
            running_future += row.get('pvFuture', 0)
            cumulative_past.append(running_past)
            cumulative_future.append(running_future)

        ax.fill_between(years, 0, cumulative_past, alpha=0.5, label='Cumulative Past', color='#ff9999')
        ax.fill_between(years, cumulative_past, [p + f for p, f in zip(cumulative_past, cumulative_future)],
                       alpha=0.5, label='Cumulative Future (PV)', color='#66b3ff')

        ax.plot(years, cumulative_past, linewidth=2, color='#d62728')
        ax.plot(years, [p + f for p, f in zip(cumulative_past, cumulative_future)], linewidth=2, color='#1f77b4')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Cumulative Damages ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3)

        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_age_progression_chart(rows, title="Economic Loss by Age"):
        """Create chart showing loss by age"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        ages = [row.get('age', '') for row in rows]
        losses = [row.get('loss', 0) for row in rows]

        ax.plot(ages, losses, marker='o', linewidth=2, color='#d62728', markersize=6)
        ax.fill_between(ages, 0, losses, alpha=0.3, color='#d62728')

        ax.set_xlabel('Age', fontsize=12, fontweight='bold')
        ax.set_ylabel('Annual Economic Loss ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_total_compensation_comparison_chart(rows, title="Total Compensation Comparison", include_legals=True):
        """Create stacked bar chart comparing total compensation packages"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(14, 7))
        years = [row.get('year', '') for row in rows]

        # But-For components
        bf_adj = [row.get('bfAdj', 0) for row in rows]
        bf_fringe = [row.get('bfFringe', 0) for row in rows]
        bf_legals = [row.get('bfLegals', 0) for row in rows]

        # Actual components
        act_e = [row.get('actE', 0) for row in rows]
        act_fringe = [row.get('actFringe', 0) for row in rows]
        act_legals = [row.get('actLegals', 0) for row in rows]

        x = np.arange(len(years))
        width = 0.35

        # But-For stack
        p1 = ax.bar(x - width/2, bf_adj, width, label='BF Earnings', color='#2ca02c', alpha=0.8)
        p2 = ax.bar(x - width/2, bf_fringe, width, bottom=bf_adj, label='BF Fringe', color='#98df8a', alpha=0.8)
        bf_bottom = [a + f for a, f in zip(bf_adj, bf_fringe)]
        if include_legals:
            p3 = ax.bar(x - width/2, bf_legals, width, bottom=bf_bottom, label='BF Legals', color='#d5e8d4', alpha=0.8)
            bf_bottom = [b + l for b, l in zip(bf_bottom, bf_legals)]

        # Actual stack
        p4 = ax.bar(x + width/2, act_e, width, label='Actual Earnings', color='#d62728', alpha=0.8)
        p5 = ax.bar(x + width/2, act_fringe, width, bottom=act_e, label='Actual Fringe', color='#ff9896', alpha=0.8)
        act_bottom = [a + f for a, f in zip(act_e, act_fringe)]
        if include_legals:
            p6 = ax.bar(x + width/2, act_legals, width, bottom=act_bottom, label='Actual Legals', color='#ffcccb', alpha=0.8)
            act_bottom = [b + l for b, l in zip(act_bottom, act_legals)]

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Total Compensation ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(years, rotation=45 if len(years) > 15 else 0, ha='right' if len(years) > 15 else 'center')
        ax.legend(loc='upper left', fontsize=9, ncol=2)
        ax.grid(True, alpha=0.3, axis='y')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_survival_probability_chart(rows, title="Survival Probability Over Time"):
        """Create chart showing survival probabilities"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        survival_probs = [(row.get('survivalProb') or 1) * 100 for row in rows]

        ax.plot(years, survival_probs, marker='o', linewidth=2, color='#1f77b4', markersize=6)
        ax.fill_between(years, 0, survival_probs, alpha=0.3, color='#1f77b4')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Survival Probability (%)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_ylim(0, 105)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=50, color='red', linestyle='--', linewidth=1, alpha=0.5, label='50% Threshold')
        ax.legend()

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_pv_discount_impact_chart(rows, title="Present Value Discount Impact"):
        """Create chart showing impact of PV discounting"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        future_raw = [row.get('futurePart', 0) for row in rows]
        pv_future = [row.get('pvFuture', 0) for row in rows]

        ax.plot(years, future_raw, marker='o', linewidth=2, label='Undiscounted Future Loss', color='#ff7f0e')
        ax.plot(years, pv_future, marker='s', linewidth=2, label='Present Value', color='#2ca02c')
        ax.fill_between(years, pv_future, future_raw, alpha=0.2, color='gray', label='Discount Amount')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Future Damages ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_past_vs_future_chart(rows, title="Past vs Future Damages Distribution"):
        """Create chart showing distribution of past and future damages"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        past_dam = [row.get('pastPart', 0) for row in rows]
        future_pv = [row.get('pvFuture', 0) for row in rows]

        width = 0.35
        x = np.arange(len(years))

        ax.bar(x - width/2, past_dam, width, label='Past Damages', color='#ff9999', edgecolor='black')
        ax.bar(x + width/2, future_pv, width, label='Future PV', color='#66b3ff', edgecolor='black')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Damages ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(years, rotation=45 if len(years) > 15 else 0, ha='right' if len(years) > 15 else 'center')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3, axis='y')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_tax_impact_chart(rows, title="Tax Impact on But-For Earnings"):
        """Create chart showing gross vs after-tax earnings"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        bf_gross = [row.get('bfGross', 0) for row in rows]
        bf_adj = [row.get('bfAdj', 0) for row in rows]
        tax_impact = [gross - adj for gross, adj in zip(bf_gross, bf_adj)]

        ax.bar(years, bf_adj, label='After-Tax Earnings', color='#2ca02c', alpha=0.8, edgecolor='black')
        ax.bar(years, tax_impact, bottom=bf_adj, label='Tax Impact', color='#d62728', alpha=0.6, edgecolor='black')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Earnings ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3, axis='y')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_legally_required_benefits_chart(rows, title="Legally Required Benefits Comparison"):
        """Create chart comparing legally required benefits"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        bf_legals = [row.get('bfLegals', 0) for row in rows]
        act_legals = [row.get('actLegals', 0) for row in rows]

        ax.plot(years, bf_legals, marker='o', linewidth=2, label='But-For Legally Req', color='#2ca02c')
        ax.plot(years, act_legals, marker='s', linewidth=2, label='Actual Legally Req', color='#d62728')
        ax.fill_between(years, bf_legals, act_legals, alpha=0.2, color='gray')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Legally Required Benefits ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_retirement_scenario_timeline_chart(scenario, title="Retirement Scenario Timeline"):
        """Create timeline chart for a specific retirement scenario"""
        rows = scenario.get('schedule', {}).get('rows', [])
        if not rows:
            return None

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10))

        years = [row.get('year', '') for row in rows]
        bf_total = [(row.get('bfAdj', 0) + row.get('bfFringe', 0) + row.get('bfLegals', 0)) for row in rows]
        act_total = [(row.get('actE', 0) + row.get('actFringe', 0) + row.get('actLegals', 0)) for row in rows]
        loss = [row.get('loss', 0) for row in rows]

        # Top chart: Earnings comparison
        ax1.plot(years, bf_total, marker='o', linewidth=2, label='But-For Total', color='#2ca02c')
        ax1.plot(years, act_total, marker='s', linewidth=2, label='Actual Total', color='#d62728')
        ax1.fill_between(years, bf_total, act_total, alpha=0.2, color='gray')
        ax1.set_ylabel('Total Compensation ($)', fontsize=11, fontweight='bold')
        ax1.set_title(f'{title} - Compensation Comparison', fontsize=13, fontweight='bold')
        ax1.legend(loc='best', fontsize=10)
        ax1.grid(True, alpha=0.3)
        ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        # Bottom chart: Annual loss
        ax2.bar(years, loss, color='#d62728', alpha=0.7, edgecolor='black')
        ax2.set_xlabel('Year', fontsize=11, fontweight='bold')
        ax2.set_ylabel('Annual Loss ($)', fontsize=11, fontweight='bold')
        ax2.set_title('Annual Economic Loss', fontsize=13, fontweight='bold')
        ax2.grid(True, alpha=0.3, axis='y')
        ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        if len(years) > 15:
            for ax in [ax1, ax2]:
                ax.tick_params(axis='x', rotation=45)
                for label in ax.get_xticklabels():
                    label.set_ha('right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_loss_percentage_chart(rows, title="Economic Loss as Percentage of But-For Earnings"):
        """Create chart showing loss percentage over time"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        loss_percentages = []

        for row in rows:
            bf_total = row.get('bfAdj', 0) + row.get('bfFringe', 0) + row.get('bfLegals', 0)
            loss = row.get('loss', 0)
            loss_pct = (loss / bf_total * 100) if bf_total > 0 else 0
            loss_percentages.append(loss_pct)

        colors = ['#d62728' if pct >= 75 else '#ff7f0e' if pct >= 50 else '#2ca02c' for pct in loss_percentages]
        ax.bar(years, loss_percentages, color=colors, alpha=0.7, edgecolor='black')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Loss Percentage (%)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3, axis='y')
        ax.axhline(y=50, color='orange', linestyle='--', linewidth=1, alpha=0.5, label='50% Loss')
        ax.axhline(y=75, color='red', linestyle='--', linewidth=1, alpha=0.5, label='75% Loss')
        ax.legend()

        if len(years) > 15:
            plt.xticks(rotation=45, ha='right')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_jury_items_chart(totals, title="What Was Lost"):
        """Simple bar for past vs future damages"""
        if totals is None:
            totals = {}
        past = totals.get('pastDam', 0) or 0
        future = totals.get('futurePV', 0) or 0
        if past <= 0 and future <= 0:
            return None
        fig, ax = plt.subplots(figsize=(6, 4))
        labels = ['Past Damages', 'Future Damages (PV)']
        vals = [past, future]
        colors = ['#ff9999', '#66b3ff']
        bars = ax.bar(labels, vals, color=colors, edgecolor='black')
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2, h, f"${h:,.0f}", ha='center', va='bottom', fontsize=10, fontweight='bold')
        ax.grid(True, axis='y', alpha=0.3)
        plt.tight_layout()
        stream = io.BytesIO()
        plt.savefig(stream, format='png', dpi=300, bbox_inches='tight')
        stream.seek(0)
        plt.close(fig)
        return stream

    def create_jury_years_chart(schedule, title="How Long The Loss Lasts"):
        """Simple bar for total years of loss"""
        rows = schedule.get('rows', []) if isinstance(schedule, dict) else []
        if not rows:
            return None
        years_of_loss = 0
        for r in rows:
            try:
                years_of_loss += float(r.get('portion', 1) or 0)
            except Exception:
                years_of_loss += 1
        if years_of_loss <= 0:
            years_of_loss = len(rows)
        fig, ax = plt.subplots(figsize=(5, 4))
        bar = ax.bar(['Years with Loss'], [years_of_loss], color='#8dd3c7', edgecolor='black')
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_ylim(0, max(years_of_loss * 1.2, 1))
        for b in bar:
            h = b.get_height()
            ax.text(b.get_x() + b.get_width()/2, h, f"{h:,.1f} years", ha='center', va='bottom', fontsize=10, fontweight='bold')
        ax.grid(True, axis='y', alpha=0.3)
        plt.tight_layout()
        stream = io.BytesIO()
        plt.savefig(stream, format='png', dpi=300, bbox_inches='tight')
        stream.seek(0)
        plt.close(fig)
        return stream

    def create_jury_growth_chart(assumptions, title="Growth / Inflation Factor"):
        """Simple bar for growth/inflation rate"""
        if not isinstance(assumptions, dict):
            return None
        but_for = assumptions.get('butFor', {}) or {}
        method = but_for.get('growthMethod')
        growth_val = but_for.get('growth', 0) or 0
        label = ''
        if method == 'fixed':
            label = f"{growth_val*100:.2f}% fixed growth"
        elif method == 'ups':
            label = 'UPS contract growth (varies by year)'
        elif method == 'series':
            label = 'Growth varies by series'
        else:
            label = f"{growth_val*100:.2f}% growth"
        fig, ax = plt.subplots(figsize=(6, 4))
        bar = ax.bar(['Growth / Inflation'], [growth_val * 100], color='#80b1d3', edgecolor='black')
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x:.2f}%'))
        ymax = max((growth_val * 100) * 1.4, 1)
        ax.set_ylim(0, ymax)
        for b in bar:
            h = b.get_height()
            ax.text(b.get_x() + b.get_width()/2, h, f"{h:.2f}%", ha='center', va='bottom', fontsize=10, fontweight='bold')
        if label:
            ax.text(0, ymax * 0.8, label, ha='center', fontsize=9, fontstyle='italic')
        ax.grid(True, axis='y', alpha=0.3)
        plt.tight_layout()
        stream = io.BytesIO()
        plt.savefig(stream, format='png', dpi=300, bbox_inches='tight')
        stream.seek(0)
        plt.close(fig)
        return stream

    def create_jury_total_chart(totals, title="Total Economic Loss"):
        """Simple bar for total PV"""
        if totals is None:
            totals = {}
        total_pv = totals.get('totalPV', 0)
        if total_pv <= 0:
            return None
        fig, ax = plt.subplots(figsize=(5, 4))
        bar = ax.bar(['Total Present Value'], [total_pv], color='#fdb462', edgecolor='black')
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        for b in bar:
            h = b.get_height()
            ax.text(b.get_x() + b.get_width()/2, h, f"${h:,.0f}", ha='center', va='bottom', fontsize=11, fontweight='bold')
        ax.grid(True, axis='y', alpha=0.3)
        plt.tight_layout()
        stream = io.BytesIO()
        plt.savefig(stream, format='png', dpi=300, bbox_inches='tight')
        stream.seek(0)
        plt.close(fig)
        return stream


    def create_ups_fringe_breakdown_chart(rows, title="UPS Fringe Benefits Components"):
        """Create detailed UPS fringe breakdown chart"""
        if not rows:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        years = [row.get('year', '') for row in rows]
        hw = [row.get('bfHW', 0) for row in rows]
        pension = [row.get('bfPension', 0) for row in rows]

        x = np.arange(len(years))
        width = 0.35

        ax.bar(x - width/2, hw, width, label='Health & Welfare', color='#8c564b', alpha=0.8, edgecolor='black')
        ax.bar(x + width/2, pension, width, label='Pension', color='#e377c2', alpha=0.8, edgecolor='black')

        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_ylabel('Fringe Benefit Amount ($)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(years, rotation=45 if len(years) > 15 else 0, ha='right' if len(years) > 15 else 'center')
        ax.legend(loc='best', fontsize=11)
        ax.grid(True, alpha=0.3, axis='y')
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def create_sensitivity_heatmap(sensitivity):
        """Create heatmap for sensitivity analysis"""
        if not sensitivity or not sensitivity.get('results'):
            return None

        disc_range = sensitivity.get('discountRange', sensitivity.get('range', []))
        growth_range = sensitivity.get('growthRange', sensitivity.get('range', []))
        results = sensitivity.get('results', [])
        base_disc = sensitivity.get('baseDiscountRate', 0)
        base_growth = sensitivity.get('baseGrowthRate', 0)

        if not disc_range or not results:
            return None

        # Build matrix
        matrix = []
        for i, disc_delta in enumerate(disc_range):
            if i < len(results):
                row_data = results[i]
                if isinstance(row_data, list):
                    row_values = [cell.get('totalPV', 0) if isinstance(cell, dict) else 0 for cell in row_data]
                    matrix.append(row_values)

        if not matrix:
            return None

        fig, ax = plt.subplots(figsize=(10, 8))

        im = ax.imshow(matrix, cmap='RdYlGn', aspect='auto')

        # Set ticks
        ax.set_xticks(np.arange(len(growth_range)))
        ax.set_yticks(np.arange(len(disc_range)))

        # Label ticks
        growth_labels = [f"{(base_growth + g)*100:.1f}%" for g in growth_range]
        disc_labels = [f"{(base_disc + d)*100:.1f}%" for d in disc_range]

        ax.set_xticklabels(growth_labels)
        ax.set_yticklabels(disc_labels)

        ax.set_xlabel('Growth Rate', fontsize=12, fontweight='bold')
        ax.set_ylabel('Discount Rate', fontsize=12, fontweight='bold')
        ax.set_title('Sensitivity Analysis: Total Present Value', fontsize=14, fontweight='bold')

        # Add text annotations
        for i in range(len(disc_range)):
            for j in range(len(growth_range)):
                if i < len(matrix) and j < len(matrix[i]):
                    text = ax.text(j, i, f'${matrix[i][j]:,.0f}',
                                 ha="center", va="center", color="black", fontsize=8, fontweight='bold')

        # Colorbar
        cbar = plt.colorbar(im, ax=ax)
        cbar.set_label('Total PV ($)', rotation=270, labelpad=20, fontweight='bold')

        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    @app.route('/api/export/word', methods=['POST'])
    def export_word():
        """Export damage schedule to Word document"""
        try:
            data = request.json
            if not data:
                return jsonify({'success': False, 'error': 'No data provided'}), 400

            # Create Word document
            doc = Document()

            # Set landscape orientation
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run('But-For Damages Analysis Report')
            title_run.bold = True
            title_run.font.size = Pt(16)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add case information
            case_name = data.get('assumptions', {}).get('meta', {}).get('caseName', 'Unnamed Case')
            doc.add_paragraph(f"Case: {case_name}")
            doc.add_paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph()

            # Add summary totals
            totals = data.get('schedule', {}).get('totals', {})
            summary = doc.add_paragraph()
            summary.add_run('Summary Totals\n').bold = True
            include_discounting = data.get('assumptions', {}).get('options', {}).get('includeDiscounting', True)
            future_nominal_total = sum(r.get('futurePart', 0) for r in data.get('schedule', {}).get('rows', []) or [])
            if include_discounting:
                summary.add_run(f"Total Present Value: ${totals.get('totalPV', 0):,.2f}\n")
                summary.add_run(f"Past Damages: ${totals.get('pastDam', 0):,.2f}\n")
                summary.add_run(f"Future PV: ${totals.get('futurePV', 0):,.2f}")
            else:
                nominal_total = (totals.get('pastDam', 0) or 0) + future_nominal_total
                summary.add_run(f"Total (Nominal): ${nominal_total:,.2f}\n")
                summary.add_run(f"Past Damages: ${totals.get('pastDam', 0):,.2f}\n")
                summary.add_run(f"Future (Nominal): ${future_nominal_total:,.2f}")
            active_scenario = data.get('activeScenario', {})
            if active_scenario.get('name'):
                summary.add_run(f"\nScenario: {active_scenario['name']}")
            doc.add_paragraph()

            assumptions = data.get('assumptions', {})
            meta = assumptions.get('meta', {})
            life_table = assumptions.get('lifeTable', {})
            validation_notes = _validate_assumption_ranges(assumptions)
            provenance = _build_provenance_metadata(assumptions)

            schedule = data.get('schedule', {})
            all_rows = schedule.get('rows', [])
            use_ups_fringe = assumptions.get('butFor', {}).get('fringeMethod') == 'ups'
            include_legals = assumptions.get('options', {}).get('includeLegals', True)
            aef_on = assumptions.get('options', {}).get('includeAEF', False) and assumptions.get('aef', {}).get('mode') == 'on'
            tinari_mode = aef_on
            show_fringe = not aef_on

            doc.add_heading('Assumptions Overview', level=1)
            doc_notes = []
            if life_table:
                lt_source = life_table.get('source')
                lt_population = life_table.get('population', 'combined').title()
                doc_notes.append(f"Life table: {lt_source or 'CDC United States Life Tables (default)'} [{lt_population}]")
            if meta.get('wageSourceNotes'):
                doc_notes.append(f"Wage/growth documentation: {meta['wageSourceNotes']}")
            if meta.get('benefitSourceNotes'):
                doc_notes.append(f"Fringe/benefit documentation: {meta['benefitSourceNotes']}")

            for note in doc_notes:
                doc.add_paragraph(note)
            if doc_notes:
                doc.add_paragraph()

            doc.add_heading('Provenance & Validation', level=2)
            doc.add_paragraph(f"Generated (UTC): {provenance['generated_at']}")
            doc.add_paragraph(f"Assumptions fingerprint (SHA-256): {provenance['fingerprint']}")
            if provenance['sources']:
                doc.add_paragraph('Sources:')
                for source in provenance['sources']:
                    doc.add_paragraph(source, style='List Bullet')
            validation_header = doc.add_paragraph()
            validation_header.add_run('Validation status: ').bold = True
            if validation_notes:
                validation_header.add_run('Attention required').font.color.rgb = RGBColor(0xB0, 0x1E, 0x1E)
                for warning in validation_notes:
                    doc.add_paragraph(warning, style='List Bullet')
            else:
                validation_header.add_run('All monitored inputs fall within configured ranges.').font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
            doc.add_paragraph()

            # ==================== AEF BREAKDOWN TABLE ====================
            doc.add_heading('Adjusted Earnings Factor (AEF) Breakdown', level=1)
            doc.add_paragraph('This table shows how the AEF is calculated from its component parts.')
            doc.add_paragraph(f"Formula: AEF = (WLE/YFS) × (1 - UR×(1-URF)) × (1 - TL_eff) × (1 + FB){' × (1 - PC) × (1 - PM)' if meta.get('caseType') == 'wd' else ''}; Adjusted Earnings = GE × AEF.")
            doc.add_paragraph()

            aef_data = assumptions.get('aef', {})
            horizon = assumptions.get('horizon', {})
            is_wd = meta.get('caseType') == 'wd'

            # Create AEF table
            aef_table = doc.add_table(rows=1, cols=3)
            aef_table.style = 'Light Grid Accent 1'

            # Set headers
            hdr_cells = aef_table.rows[0].cells
            hdr_cells[0].text = 'Component'
            hdr_cells[1].text = 'Value'
            hdr_cells[2].text = 'Description'

            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Helper function to add AEF row
            def add_aef_row(component, value, description, is_header=False):
                row = aef_table.add_row()
                cells = row.cells
                cells[0].text = component
                cells[1].text = value
                cells[2].text = description

                if is_header:
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(11)

            # Base Components
            add_aef_row('Base Components', '', '', True)
            but_for_data = assumptions.get('butFor', {})
            gross_earnings_base = aef_data.get('grossEarningsBase', but_for_data.get('salary', 0))
            add_aef_row('Gross Earnings Base (GE)', f"${gross_earnings_base:,.2f}",
                       'Base annual but-for salary')

            # Work-Life & Unemployment Components
            add_aef_row('Work-Life & Unemployment Components', '', '', True)
            add_aef_row('Worklife Adjusted Earnings Base (WLE)', f"{aef_data.get('wle', 1):.5f}",
                       f"WLE / YFS = {horizon.get('wleYears', 0)} / {horizon.get('yfsYears', 0)}")
            add_aef_row('Unemployment Rate (UR)', f"{aef_data.get('UR', 0)*100:.2f}%", 'Probability of unemployment')
            add_aef_row('Unemployment Reimbursement (URF)', f"{aef_data.get('URF', 0)*100:.2f}%",
                       'Portion of unemployment offset by benefits')
            add_aef_row('Effective Unemployment', f"{aef_data.get('ufEff', 0)*100:.2f}%",
                       f"UR × (1 - URF) = {aef_data.get('UR', 0)*100:.2f}% × {(1 - aef_data.get('URF', 0)):.3f}")

            # Tax Components
            add_aef_row('Tax Components', '', '', True)
            add_aef_row('Federal Tax (TLF)', f"{aef_data.get('TLF', 0)*100:.2f}%", 'Federal income tax rate')
            add_aef_row('State Tax (TLS)', f"{aef_data.get('TLS', 0)*100:.2f}%", 'State income tax rate')
            add_aef_row('Combined Tax (Simple Add)', f"{aef_data.get('tlCombined', 0)*100:.2f}%",
                       'TLF + TLS (for reference only)')
            add_aef_row('Effective Tax Rate', f"{aef_data.get('tlEff', 0)*100:.2f}%",
                       '1 - (1 - TLF) × (1 - TLS) = Multiplicative method')

            # Wrongful Death Components (if applicable)
            if is_wd:
                add_aef_row('Wrongful Death Components', '', '', True)
                add_aef_row('Personal Consumption (PC)', f"{aef_data.get('PC', 0)*100:.2f}%",
                           "Decedent's personal consumption")
                add_aef_row('Personal Maintenance (PM)', f"{aef_data.get('PM', 0)*100:.2f}%",
                           "Decedent's personal maintenance")

            # Fringe Benefit Load (built into AEF)
            add_aef_row('Fringe Benefit Load', '', '', True)
            fringe_pct = aef_data.get('fringePct', 0)
            add_aef_row('Fringe Benefit %', f"{fringe_pct*100:.2f}%",
                       'Fringe benefits as % of wages')
            add_aef_row('Fringe Load Factor', f"{(1 + fringe_pct):.5f}",
                       '(1 + Fringe %) - multiplier applied in AEF')

            # Final AEF Calculation
            add_aef_row('Adjusted Earnings Factor (AEF)', '', '', True)
            formula = 'GE × WLE × (1 - UF) × (1 - TR) × (1 + FB)'
            if is_wd:
                formula += ' × (1 - PC) × (1 - PM)'

            factor_formula = 'WLE × (1 - UF) × (1 - TR) × (1 + FB)'
            if is_wd:
                factor_formula += ' × (1 - PC) × (1 - PM)'

            add_aef_row('AEF Factor',
                       f"{aef_data.get('factor', 1):.5f}" if aef_on else '1.00000 (OFF)',
                       factor_formula)

            aef_value = aef_data.get('value', 0)
            add_aef_row('AEF (Adjusted Annual Earnings)',
                       f"${aef_value:,.2f}" if aef_on else 'N/A (OFF)',
                       formula)

            add_aef_row('AEF Applied To', 'Gross Wages' if aef_on else 'N/A',
                       'Result = Wages × AEF Factor (includes wage + fringe adjustment)' if aef_on else 'AEF is turned off')
            add_aef_row('Fringe Treatment',
                       'Built into AEF' if aef_on else 'Added separately',
                       'Fringe benefits included via (1 + FB) multiplier in AEF' if aef_on else 'Fringes calculated and added separately when AEF is off')

            doc.add_paragraph()

            def add_formula_note(text):
                note = doc.add_paragraph(f"Formula: {text}")
                for run in note.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                return note

            def has_nonzero(rows, key):
                """Return True if any row has a non-zero/meaningful value for key."""
                for r in rows or []:
                    try:
                        val = r.get(key, 0)
                        if isinstance(val, str):
                            if val.strip() == '':
                                continue
                            cleaned = val.replace('$', '').replace(',', '')
                            val = float(cleaned)
                        if isinstance(val, (int, float)) and abs(val) > 1e-9:
                            return True
                    except Exception:
                        continue
                return False

            def add_schedule_table(title, rows, use_ups_fringe, aef_on=False, include_legals=True, include_discounting=True, description=None):
                doc.add_heading(title, level=2)

                if description:
                    desc_para = doc.add_paragraph(description)
                    desc_para.style = 'Normal'
                    for run in desc_para.runs:
                        run.font.italic = True
                        run.font.size = Pt(10)

                add_formula_note("Loss = BF After-Tax (+Fringe if shown) (+Legals if shown) - (ACT Earn + ACT Fringe (+ACT Legals if shown)); Past shown in 'Past' column; Future columns include PV only when discounting is on.")

                headers = ['Year', 'Age', 'Portion', 'BF Gross', 'BF After-Tax']
                hw_used = pension_used = fringe_total_used = False
                fringe_used = False
                if not aef_on:
                    if use_ups_fringe:
                        hw_used = has_nonzero(rows, 'bfHW')
                        pension_used = has_nonzero(rows, 'bfPension')
                        fringe_total_used = has_nonzero(rows, 'bfFringe')
                        if hw_used:
                            headers.append('H&W')
                            fringe_used = True
                        if pension_used:
                            headers.append('Pension')
                            fringe_used = True
                        if fringe_total_used:
                            headers.append('Tot Fringe')
                            fringe_used = True
                    else:
                        fringe_used = has_nonzero(rows, 'bfFringe')
                        if fringe_used:
                            headers.append('BF Fringe')
                legals_used = include_legals and has_nonzero(rows, 'bfLegals')
                if legals_used:
                    headers.append('BF Legals')

                act_fringe_used = has_nonzero(rows, 'actFringe')
                act_legals_used = include_legals and has_nonzero(rows, 'actLegals')

                headers.append('ACT Earn')
                if act_fringe_used:
                    headers.append('ACT Fringe')
                if act_legals_used:
                    headers.append('ACT Legals')
                headers.extend(['Loss', 'Past'])
                if include_discounting:
                    headers.extend(['Future Raw', 'Future Surv', 'PV Future', 'Surv Prob'])
                else:
                    headers.append('Future')

                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                table.allow_autofit = False

                col_width = Inches(10.0 / len(headers))
                for col in table.columns:
                    col.width = col_width

                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].width = col_width
                    for paragraph in hdr_cells[i].paragraphs:
                        paragraph.paragraph_format.keep_together = True
                        paragraph.paragraph_format.widow_control = True
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(8)

                for row_data in rows:
                    row_cells = table.add_row().cells
                    values = [
                        str(row_data['year']),
                        row_data['age'],
                        f"{row_data.get('portion', 0):.3f}" if isinstance(row_data.get('portion'), (int, float)) else str(row_data.get('portion', '')),
                        f"${row_data.get('bfGross', 0):,.2f}",
                        f"${row_data.get('bfAdj', 0):,.2f}"
                    ]

                    if not aef_on:
                        if use_ups_fringe:
                            if hw_used:
                                values.append(f"${row_data.get('bfHW', 0):,.2f}")
                            if pension_used:
                                values.append(f"${row_data.get('bfPension', 0):,.2f}")
                            if fringe_total_used:
                                values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                        else:
                            if fringe_used:
                                values.append(f"${row_data.get('bfFringe', 0):,.2f}")

                    if legals_used:
                        values.append(f"${row_data.get('bfLegals', 0):,.2f}")

                    values.append(f"${row_data.get('actE', 0):,.2f}")
                    if act_fringe_used:
                        values.append(f"${row_data.get('actFringe', 0):,.2f}")
                    if act_legals_used:
                        values.append(f"${row_data.get('actLegals', 0):,.2f}")
                    values.extend([
                        f"${row_data.get('loss', 0):,.2f}",
                        f"${row_data.get('pastPart', 0):,.2f}",
                    ])
                    if include_discounting:
                        values.extend([
                            f"${row_data.get('futurePart', 0):,.2f}",
                            f"${row_data.get('survivalWeightedFuture', row_data.get('futurePart', 0)):,.2f}",
                            f"${row_data.get('pvFuture', 0):,.2f}",
                            f"{(row_data.get('survivalProb') or 1):,.3f}"
                        ])
                    else:
                        values.append(f"${row_data.get('futurePart', 0):,.2f}")

                    for idx, value in enumerate(values):
                        if idx < len(row_cells):
                            row_cells[idx].text = value
                            row_cells[idx].width = col_width
                            for paragraph in row_cells[idx].paragraphs:
                                paragraph.paragraph_format.keep_together = True
                                for run in paragraph.runs:
                                    run.font.size = Pt(7)
                                    run.font.name = 'Calibri'

                doc.add_paragraph()

            def add_tinari_table(rows_pre, rows_post, assumptions):
                doc.add_heading('TINARI-STYLE SUMMARY (AEF ON)', level=2)
                doc.add_paragraph('Formatted to match Tinari presentation (Year, Age, Portion, Base Earnings, Adjusted Income, Present Value).')

                headers = ['Year', 'Age', 'Portion of Year', 'Base Earnings', 'Adjusted Income', 'Present Value']
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                for i, h in enumerate(headers):
                    hdr = table.rows[0].cells[i]
                    hdr.text = h
                    for run in hdr.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

                aef = assumptions.get('aef', {})
                options = assumptions.get('options', {})
                discount = assumptions.get('discount', {})
                discount_rate = (discount.get('ndr') if discount.get('method') == 'ndr' else discount.get('rate', 0)) or 0

                def pct(v):
                    return f"{v*100:.2f}%" if v is not None else ''

                formula_parts = []
                wle = aef.get('wle'); uf = aef.get('ufEff'); tl = aef.get('tlEff'); fb = aef.get('fringePct')
                formula_parts.append(f"(WLE/YFS {wle:.5f})" if isinstance(wle, (int, float)) else "(WLE/YFS)")
                formula_parts.append(f"× (1 - UR×(1-URF) {uf:.5f})" if isinstance(uf, (int, float)) else "× (1 - UR×(1-URF))")
                formula_parts.append(f"× (1 - TL_eff {tl:.5f})" if isinstance(tl, (int, float)) else "× (1 - TL_eff)")
                formula_parts.append(f"× (1 + FB {fb:.5f})" if isinstance(fb, (int, float)) else "× (1 + FB)")
                if assumptions.get('meta', {}).get('caseType') == 'wd':
                    pc = aef.get('PC', 0); pm = aef.get('PM', 0)
                    formula_parts.append(f"× (1 - PC {pc:.5f})" if isinstance(pc, (int, float)) else "× (1 - PC)")
                    formula_parts.append(f"× (1 - PM {pm:.5f})" if isinstance(pm, (int, float)) else "× (1 - PM)")
                factor = aef.get('factor')
                if isinstance(factor, (int, float)):
                    formula_parts.append(f"= {factor:.5f}")
                formula_str = ' '.join(formula_parts)

                formula_row = table.add_row().cells
                formula_row[0].text = ''
                formula_row[1].text = ''
                growth_label = ''
                but_for = assumptions.get('butFor', {})
                if but_for.get('growthMethod') == 'fixed':
                    g = but_for.get('growth', 0)
                    growth_label = pct(g) + " growth"
                elif but_for.get('growthMethod') == 'ups':
                    growth_label = "UPS contract growth (varies)"
                elif but_for.get('growthMethod') == 'series':
                    growth_label = "Series growth (varies)"

                formula_row[2].text = growth_label
                formula_row[3].text = formula_str
                formula_row[4].text = f"{pct(discount_rate)} discount"
                formula_row[5].text = ''
                for cell in formula_row:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)
                        run.font.italic = True

                eq_row = table.add_row().cells
                eq_row[0].text = ''
                eq_row[1].text = ''
                eq_row[2].text = ''
                eq_row[3].text = 'Adjusted Income = Base × AEF'
                eq_row[4].text = 'Present Value = Adjusted / (1 + r)^{years from valuation}'
                eq_row[5].text = ''
                for cell in eq_row:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)
                        run.font.italic = True

                def add_section(label):
                    r = table.add_row().cells
                    r[0].text = label
                    for p in r[0].paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                    for i in range(1, len(r)):
                        r[i].text = ''

                def add_data(rows, label):
                    adj_total = 0
                    pv_total = 0

                    for row in rows:
                        cells = table.add_row().cells
                        pv = (row.get('pastPart', 0) or 0) + (row.get('pvFuture', 0) or 0)
                        adj = row.get('bfAdj', 0) or 0
                        adj_total += adj
                        pv_total += pv

                        values = [
                            str(row.get('year', '')),
                            str(row.get('age', '')),
                            f"{(row.get('portion', 0)*100):.0f}%",
                            f"${row.get('bfGross', 0):,.2f}",
                            f"${adj:,.2f}",
                            f"${pv or 0:,.2f}"
                        ]
                        for i, val in enumerate(values):
                            cells[i].text = val
                            for run in cells[i].paragraphs[0].runs:
                                run.font.size = Pt(8)

                    if rows:
                        total_cells = table.add_row().cells
                        total_cells[0].text = f"{label} Totals"
                        total_cells[1].text = ''
                        total_cells[2].text = ''
                        total_cells[3].text = ''
                        total_cells[4].text = f"${adj_total:,.2f}"
                        total_cells[5].text = f"${pv_total:,.2f}"
                        for idx in (0, 4, 5):
                            for run in total_cells[idx].paragraphs[0].runs:
                                run.font.size = Pt(8)
                                run.font.bold = True

                    return adj_total, pv_total

                add_section('Past Years')
                past_adj_total, past_pv_total = add_data(rows_pre or [], 'Past')
                add_section('Future Years')
                future_adj_total, future_pv_total = add_data(rows_post or [], 'Future')

                total_row = table.add_row().cells
                total_row[0].text = 'Total'
                total_row[1].text = ''
                total_row[2].text = ''
                total_row[3].text = ''
                total_adj = past_adj_total + future_adj_total
                totals_from_assumps = assumptions.get('totals') if isinstance(assumptions.get('totals'), dict) else {}
                total_pv = totals_from_assumps.get('totalPV')
                if total_pv is None:
                    total_pv = past_pv_total + future_pv_total

                total_row[4].text = f"${total_adj:,.2f}"
                total_row[5].text = f"${total_pv:,.2f}"
                for idx in (0, 4, 5):
                    for run in total_row[idx].paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

                doc.add_paragraph()

            def add_yoy_summary_table(rows, totals_context, use_ups_fringe, show_fringe, include_legals, heading, description=None, include_discounting=True):
                """Add a condensed year-over-year summary table with zero-only columns hidden."""
                if not rows:
                    return

                doc.add_heading(heading, level=2)
                if description:
                    para = doc.add_paragraph(description)
                    for run in para.runs:
                        run.italic = True

                add_formula_note("Loss = BF After-Tax (+Fringe if shown) (+Legals if shown) - (ACT Earn + ACT Fringe (+ACT Legals if shown)); Future values reflect PV only when discounting is enabled.")

                fringe_used = False
                hw_used = pension_used = fringe_total_used = False
                if show_fringe:
                    if use_ups_fringe:
                        hw_used = has_nonzero(rows, 'bfHW')
                        pension_used = has_nonzero(rows, 'bfPension')
                        fringe_total_used = has_nonzero(rows, 'bfFringe')
                        fringe_used = hw_used or pension_used or fringe_total_used
                    else:
                        fringe_used = has_nonzero(rows, 'bfFringe')
                legals_used = include_legals and has_nonzero(rows, 'bfLegals')
                act_fringe_used = has_nonzero(rows, 'actFringe')
                act_legals_used = include_legals and has_nonzero(rows, 'actLegals')

                headers = ['Year', 'Age', 'BF Gross']
                if fringe_used:
                    if use_ups_fringe:
                        if hw_used:
                            headers.append('H&W')
                        if pension_used:
                            headers.append('Pension')
                        if fringe_total_used:
                            headers.append('Fringe Total')
                    else:
                        headers.append('BF Fringe')
                if legals_used:
                    headers.append('BF Legals')
                headers.append('ACT Earn')
                if act_fringe_used:
                    headers.append('ACT Fringe')
                if act_legals_used:
                    headers.append('ACT Legals')
                headers.extend(['Loss', 'Past'])
                if include_discounting:
                    headers.append('Future PV')
                else:
                    headers.append('Future (Nominal)')

                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                table.allow_autofit = False
                col_width = Inches(10.0 / len(headers))
                for col in table.columns:
                    col.width = col_width

                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].width = col_width
                    for paragraph in hdr_cells[i].paragraphs:
                        paragraph.paragraph_format.keep_together = True
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(8)

                for row_data in rows:
                    row_cells = table.add_row().cells
                    values = [
                        str(row_data.get('year', '')),
                        str(row_data.get('age', '')),
                        f"${row_data.get('bfGross', 0):,.2f}"
                    ]
                    if fringe_used:
                        if use_ups_fringe:
                            if hw_used:
                                values.append(f"${row_data.get('bfHW', 0):,.2f}")
                            if pension_used:
                                values.append(f"${row_data.get('bfPension', 0):,.2f}")
                            if fringe_total_used:
                                values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                        else:
                            values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                    if legals_used:
                        values.append(f"${row_data.get('bfLegals', 0):,.2f}")

                    values.append(f"${row_data.get('actE', 0):,.2f}")
                    if act_fringe_used:
                        values.append(f"${row_data.get('actFringe', 0):,.2f}")
                    if act_legals_used:
                        values.append(f"${row_data.get('actLegals', 0):,.2f}")
                values.extend([
                    f"${row_data.get('loss', 0):,.2f}",
                    f"${row_data.get('pastPart', 0):,.2f}",
                ])
                if include_discounting:
                    values.append(f"${row_data.get('pvFuture', 0):,.2f}")
                else:
                    values.append(f"${row_data.get('futurePart', 0):,.2f}")

                    for idx, value in enumerate(values):
                        if idx < len(row_cells):
                            row_cells[idx].text = value
                            row_cells[idx].width = col_width
                            for paragraph in row_cells[idx].paragraphs:
                                paragraph.paragraph_format.keep_together = True
                                for run in paragraph.runs:
                                    run.font.size = Pt(7)

                total_row = table.add_row().cells
                total_values = []
                future_label = 'Future PV' if include_discounting else 'Future (Nominal)'
                for header in headers:
                    if header == 'Year':
                        total_values.append('TOTAL')
                    elif header == 'Past':
                        total_values.append(f"${totals_context.get('pastDam', 0):,.2f}")
                    elif header == future_label:
                        if include_discounting:
                            total_values.append(f"${totals_context.get('futurePV', 0):,.2f}")
                        else:
                            total_values.append(f"${sum(r.get('futurePart', 0) for r in rows or []):,.2f}")
                    else:
                        total_values.append('')
                for idx, value in enumerate(total_values):
                    total_row[idx].text = value
                    total_row[idx].width = col_width
                    for paragraph in total_row[idx].paragraphs:
                        paragraph.paragraph_format.keep_together = True
                        for run in paragraph.runs:
                            run.font.size = Pt(7)
                            if idx in [0, headers.index('Past'), headers.index(future_label)]:
                                run.font.bold = True

                qa_sum = sum(r.get('pvFuture', r.get('futurePart', 0)) for r in rows or [])
                qa_diff = abs(qa_sum - totals_context.get('futurePV', qa_sum))
                qa_para = doc.add_paragraph()
                if qa_diff <= 1:
                    qa_para.add_run("QA ✓ Year-over-year PV matches scenario Future PV").bold = True
                else:
                    qa_para.add_run(f"QA Warning: Sum of YOY PV (${qa_sum:,.2f}) differs from scenario Future PV (${totals_context.get('futurePV', 0):,.2f}).").italic = True
                doc.add_paragraph()
            # Add tables with clear labels and descriptions in a simplified order
            doc.add_heading('Tinari Tables', level=1)
            if tinari_mode:
                assumptions_for_tinari = data.get('assumptions', {}).copy()
                assumptions_for_tinari['totals'] = totals
                add_tinari_table(schedule.get('rowsPre', []), schedule.get('rowsPost', []), assumptions_for_tinari)
            else:
                doc.add_paragraph('Tinari layout tables appear when AEF is ON. Enable AEF to view the Tinari summary tables.')
            doc.add_paragraph()

            # Add Retirement Scenarios table
            # Add Retirement Scenarios tables and visuals
            retirement_scenarios = data.get('retirementScenarios', [])
            if retirement_scenarios and isinstance(retirement_scenarios, list) and len(retirement_scenarios) > 0:
                try:
                    # Summary (Tinari text or full table)
                    if tinari_mode:
                        doc.add_heading('Retirement Age Scenarios (Tinari Visuals)', level=2)
                        doc.add_paragraph('Tinari mode shows chart-driven retirement scenarios. Totals by age are summarized below; detailed tables follow later.').italic = True
                        add_formula_note("Total PV = Past Damages + Future PV; Future PV is computed by discounting and survival-weighting the year-by-year losses under each scenario.")
                        for scenario in retirement_scenarios:
                            if not isinstance(scenario, dict):
                                continue
                            totals_sc = scenario.get('totals', {}) if isinstance(scenario.get('totals'), dict) else {}
                            summary_para = doc.add_paragraph()
                            summary_para.add_run(f"{scenario.get('name', '') or 'Scenario'} (Age {scenario.get('retireAge', '')}): ").bold = True
                            summary_para.add_run(f"Past ${totals_sc.get('pastDam', 0):,.2f} | Future PV ${totals_sc.get('futurePV', 0):,.2f} | Total PV ${totals_sc.get('totalPV', 0):,.2f}")
                        doc.add_paragraph()
                    else:
                        doc.add_heading('TABLE 3: Retirement Age Scenarios Summary', level=2)
                        doc.add_paragraph('Comparison of total damages under different retirement age assumptions').italic = True
                        add_formula_note("Total PV = Past Damages + Future PV; each component is built from the year-by-year loss schedule for that retirement age.")
                        ret_table = doc.add_table(rows=1, cols=6)
                        ret_table.style = 'Light Grid Accent 1'
                        ret_table.allow_autofit = False
                        ret_headers = ['Scenario', 'Retire Age', 'Retire Date', 'Past Damages', 'Future PV', 'Total PV']
                        col_widths = [Inches(1.8), Inches(1.0), Inches(1.2), Inches(1.5), Inches(1.5), Inches(1.5)]
                        for idx, col in enumerate(ret_table.columns):
                            col.width = col_widths[idx]
                        hdr_cells = ret_table.rows[0].cells
                        for i, header in enumerate(ret_headers):
                            hdr_cells[i].text = header
                            hdr_cells[i].width = col_widths[i]
                            for paragraph in hdr_cells[i].paragraphs:
                                paragraph.paragraph_format.keep_together = True
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(9)
                        for scenario in retirement_scenarios:
                            if not isinstance(scenario, dict):
                                continue
                            row_cells = ret_table.add_row().cells
                            row_cells[0].text = str(scenario.get('name', ''))
                            row_cells[1].text = str(scenario.get('retireAge', ''))
                            retire_date = scenario.get('retireDate', '')
                            row_cells[2].text = retire_date[:10] if isinstance(retire_date, str) and len(retire_date) >= 10 else (str(retire_date) if retire_date else '')
                            totals = scenario.get('totals', {})
                            row_cells[3].text = f"${totals.get('pastDam', 0):,.2f}"
                            row_cells[4].text = f"${totals.get('futurePV', 0):,.2f}"
                            row_cells[5].text = f"${totals.get('totalPV', 0):,.2f}"
                            for idx, cell in enumerate(row_cells):
                                cell.width = col_widths[idx]
                                for paragraph in cell.paragraphs:
                                    paragraph.paragraph_format.keep_together = True
                                    for run in paragraph.runs:
                                        run.font.size = Pt(8)
                        doc.add_paragraph()

                    # Comparison chart
                    try:
                        ret_chart = create_retirement_scenarios_chart(retirement_scenarios, "Retirement Age Scenarios Comparison")
                        if ret_chart:
                            doc.add_heading('Retirement Scenarios Visual Comparison', level=2)
                            doc.add_paragraph('This chart compares past damages and future damages (PV) across different retirement age scenarios.')
                            doc.add_picture(ret_chart, width=Inches(7))
                            doc.add_paragraph()
                    except Exception as e:
                        print(f"Error creating retirement scenarios chart: {e}")

                    # Timeline charts (before tables)
                    doc.add_heading('Individual Retirement Scenario Analyses', level=2)
                    doc.add_paragraph("Charts for each retirement age scenario; detailed tables follow below.")
                    doc.add_paragraph()
                    for idx, scenario in enumerate(retirement_scenarios):
                        if not isinstance(scenario, dict):
                            continue
                        scenario_name = scenario.get('name', f'Scenario {idx+1}')
                        try:
                            timeline_chart = create_retirement_scenario_timeline_chart(scenario, scenario_name)
                            if timeline_chart:
                                doc.add_heading(f'{scenario_name} - Visual Timeline', level=3)
                                retire_age = scenario.get('retireAge', '')
                                doc.add_paragraph(f'Timeline visualization for retirement at age {retire_age}. Top chart shows compensation comparison, bottom chart shows annual losses.')
                                doc.add_picture(timeline_chart, width=Inches(8))
                                doc.add_paragraph()
                        except Exception as e:
                            print(f"Error creating timeline chart for {scenario_name}: {e}")

                    # Detailed year-by-year tables (always shown)
                    doc.add_heading('Retirement Scenario Year-Over-Year Tables', level=2)
                    doc.add_paragraph('Line-by-line loss tables for every retirement age scenario so jurors can follow how the totals are built.').italic = True
                    doc.add_paragraph()
                    for scenario in retirement_scenarios:
                        if not isinstance(scenario, dict):
                            continue
                        scenario_name = scenario.get('name', 'Unknown')
                        retire_age = scenario.get('retireAge', '')
                        retire_date = scenario.get('retireDate', '')
                        scenario_schedule = scenario.get('schedule', {})
                        if not scenario_schedule or not scenario_schedule.get('rows'):
                            continue
                        doc.add_heading(f"Scenario: {scenario_name} (Retire at Age {retire_age})", level=3)
                        totals_ret = scenario.get('totals', {})
                        totals_para = doc.add_paragraph()
                        totals_para.add_run(f"Retirement Date: {retire_date[:10] if isinstance(retire_date, str) and len(retire_date) >= 10 else retire_date}  |  ")
                        future_nominal_ret = sum(r.get('futurePart', 0) for r in scenario_schedule.get('rows', []) or [])
                        if include_discounting:
                            totals_para.add_run(f"Total PV: ${totals_ret.get('totalPV', 0):,.2f}  |  ")
                            totals_para.add_run(f"Past: ${totals_ret.get('pastDam', 0):,.2f}  |  ")
                            totals_para.add_run(f"Future PV: ${totals_ret.get('futurePV', 0):,.2f}")
                        else:
                            totals_para.add_run(f"Total (Nominal): ${totals_ret.get('pastDam', 0) + future_nominal_ret:,.2f}  |  ")
                            totals_para.add_run(f"Past: ${totals_ret.get('pastDam', 0):,.2f}  |  ")
                            totals_para.add_run(f"Future (Nominal): ${future_nominal_ret:,.2f}")
                        totals_para.style = 'Normal'

                        ret_options = scenario.get('assumptions', {}).get('options', data.get('assumptions', {}).get('options', {}))
                        ret_aef = scenario.get('assumptions', {}).get('aef', data.get('assumptions', {}).get('aef', {}))
                        use_ups_fringe_ret = scenario.get('assumptions', {}).get('butFor', {}).get('fringeMethod') == 'ups' or data.get('assumptions', {}).get('butFor', {}).get('fringeMethod') == 'ups'
                        aef_on_ret = ret_options.get('includeAEF', False) and ret_aef.get('mode') == 'on'
                        include_legals_ret = ret_options.get('includeLegals', True)
                        show_fringe_ret = not aef_on_ret

                        add_formula_note("Loss = BF After-Tax (+Fringe if shown) (+Legals if shown) - (ACT Earn + ACT Fringe (+ACT Legals if shown)); Past column is past portion; Future columns are precomputed per year.")

                        hw_used_ret = pension_used_ret = fringe_total_used_ret = False
                        ret_fringe_used = False
                        if show_fringe_ret:
                            if use_ups_fringe_ret:
                                hw_used_ret = has_nonzero(scenario_schedule.get('rows', []), 'bfHW')
                                pension_used_ret = has_nonzero(scenario_schedule.get('rows', []), 'bfPension')
                                fringe_total_used_ret = has_nonzero(scenario_schedule.get('rows', []), 'bfFringe')
                                ret_fringe_used = hw_used_ret or pension_used_ret or fringe_total_used_ret
                            else:
                                ret_fringe_used = has_nonzero(scenario_schedule.get('rows', []), 'bfFringe')
                        legals_used_ret = include_legals_ret and has_nonzero(scenario_schedule.get('rows', []), 'bfLegals')
                        act_fringe_used_ret = has_nonzero(scenario_schedule.get('rows', []), 'actFringe')
                        act_legals_used_ret = include_legals_ret and has_nonzero(scenario_schedule.get('rows', []), 'actLegals')

                        ret_detail_headers = ['Year', 'Age', 'Portion', 'BF Gross', 'BF After-Tax']
                        if ret_fringe_used:
                            if use_ups_fringe_ret:
                                if hw_used_ret:
                                    ret_detail_headers.append('H&W')
                                if pension_used_ret:
                                    ret_detail_headers.append('Pension')
                                if fringe_total_used_ret:
                                    ret_detail_headers.append('Tot Fringe')
                            else:
                                ret_detail_headers.append('BF Fringe')
                        if legals_used_ret:
                            ret_detail_headers.append('BF Legals')
                        ret_detail_headers.append('ACT Earn')
                        if act_fringe_used_ret:
                            ret_detail_headers.append('ACT Fringe')
                        if act_legals_used_ret:
                            ret_detail_headers.append('ACT Legals')
                        ret_detail_headers.extend(['Loss', 'Past'])
                        if include_discounting:
                            ret_detail_headers.extend(['Future Raw', 'Future Surv', 'PV Future', 'Surv Prob'])
                        else:
                            ret_detail_headers.append('Future')

                        detail_table = doc.add_table(rows=1, cols=len(ret_detail_headers))
                        detail_table.style = 'Light Grid Accent 1'
                        detail_table.allow_autofit = False
                        ret_col_width = Inches(10.0 / len(ret_detail_headers))
                        for col in detail_table.columns:
                            col.width = ret_col_width

                        hdr_cells = detail_table.rows[0].cells
                        for idx, header in enumerate(ret_detail_headers):
                            hdr_cells[idx].text = header
                            hdr_cells[idx].width = ret_col_width
                            for paragraph in hdr_cells[idx].paragraphs:
                                paragraph.paragraph_format.keep_together = True
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(7)

                        for row_data in scenario_schedule.get('rows', []):
                            if not isinstance(row_data, dict):
                                continue
                            row_cells = detail_table.add_row().cells
                            values = [
                                str(row_data.get('year', '')),
                                str(row_data.get('age', '')),
                                f"{row_data.get('portion', 0):.3f}" if isinstance(row_data.get('portion'), (int, float)) else str(row_data.get('portion', '')),
                                f"${row_data.get('bfGross', 0):,.2f}",
                                f"${row_data.get('bfAdj', 0):,.2f}"
                            ]
                            if ret_fringe_used:
                                if use_ups_fringe_ret:
                                    if hw_used_ret:
                                        values.append(f"${row_data.get('bfHW', 0):,.2f}")
                                    if pension_used_ret:
                                        values.append(f"${row_data.get('bfPension', 0):,.2f}")
                                    if fringe_total_used_ret:
                                        values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                                else:
                                    values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                            if legals_used_ret:
                                values.append(f"${row_data.get('bfLegals', 0):,.2f}")
                            values.append(f"${row_data.get('actE', 0):,.2f}")
                            if act_fringe_used_ret:
                                values.append(f"${row_data.get('actFringe', 0):,.2f}")
                            if act_legals_used_ret:
                                values.append(f"${row_data.get('actLegals', 0):,.2f}")
                            values.extend([
                                f"${row_data.get('loss', 0):,.2f}",
                                f"${row_data.get('pastPart', 0):,.2f}"
                            ])
                            if include_discounting:
                                values.extend([
                                    f"${row_data.get('futurePart', 0):,.2f}",
                                    f"${row_data.get('survivalWeightedFuture', row_data.get('futurePart', 0)):,.2f}",
                                    f"${row_data.get('pvFuture', 0):,.2f}",
                                    f"{(row_data.get('survivalProb') or 1):,.3f}"
                                ])
                            else:
                                values.append(f"${row_data.get('futurePart', 0):,.2f}")
                            for idx, value in enumerate(values):
                                if idx < len(row_cells):
                                    row_cells[idx].text = value
                                    row_cells[idx].width = ret_col_width
                                    for paragraph in row_cells[idx].paragraphs:
                                        paragraph.paragraph_format.keep_together = True
                                        for run in paragraph.runs:
                                            run.font.size = Pt(6)
                                            run.font.name = 'Calibri'
                        doc.add_paragraph()  # Spacing between scenarios

                        # Add YOY summary for this scenario
                        add_yoy_summary_table(
                            scenario_schedule.get('rows', []),
                            totals_ret,
                            use_ups_fringe_ret,
                            show_fringe_ret,
                            include_legals_ret,
                            f"Year-Over-Year Summary: {scenario_name}",
                            'Condensed summary of this retirement scenario with unused columns hidden.',
                            include_discounting
                        )
                except Exception as e:
                    print(f"Error adding retirement scenarios: {e}")
                    import traceback
                    traceback.print_exc()
                    # Continue with export even if this section fails
            # Add Sensitivity Analysis matrix
            sensitivity = data.get('sensitivityAnalysis', {})
            if include_discounting and sensitivity and sensitivity.get('results'):
                try:
                    doc.add_heading('Sensitivity Analysis', level=2)
                    doc.add_paragraph(f"Discount Method: {sensitivity.get('method', '')}")
                    doc.add_paragraph(f"Base Discount Rate: {sensitivity.get('baseDiscountRate', 0)*100:.2f}%")
                    doc.add_paragraph(f"Base Growth Rate: {sensitivity.get('baseGrowthRate', 0)*100:.2f}%")
                    add_formula_note("Each cell = Total PV computed with discount/growth deltas; PV = Σ[ (Future Loss × Survival Prob) / (1 + r)^t ] with growth applied per scenario.")
                    if tinari_mode:
                        doc.add_paragraph("Tinari mode uses the full sensitivity grid; detailed per-cell tables are suppressed for a chart-first export.").italic = True

                    disc_range = sensitivity.get('discountRange', sensitivity.get('range', []))
                    growth_range = sensitivity.get('growthRange', sensitivity.get('range', []))
                    results = sensitivity.get('results', [])
                    base_disc = sensitivity.get('baseDiscountRate', 0)
                    base_growth = sensitivity.get('baseGrowthRate', 0)
                    if tinari_mode and disc_range:
                        disc_labels = ", ".join([f"{(base_disc + d)*100:.1f}%" for d in disc_range])
                        doc.add_paragraph(f"Discount grid: {disc_labels}")
                    if tinari_mode and growth_range:
                        growth_labels = ", ".join([f"{(base_growth + g)*100:.1f}%" for g in growth_range])
                        doc.add_paragraph(f"Growth grid: {growth_labels}")
                    doc.add_paragraph()

                    if disc_range and results and len(results) > 0:
                        sens_table = doc.add_table(rows=len(disc_range)+1, cols=len(growth_range)+1)
                        sens_table.style = 'Light Grid Accent 1'

                        # Header row (growth rates)
                        sens_table.rows[0].cells[0].text = 'Disc\\Growth'
                        for j, growth_delta in enumerate(growth_range):
                            if j+1 < len(sens_table.rows[0].cells):
                                if len(growth_range) == 1 and growth_delta == 0:
                                    sens_table.rows[0].cells[j+1].text = 'Growth disabled'
                                else:
                                    growth_rate = (sensitivity.get('baseGrowthRate', 0) + growth_delta) * 100
                                    sens_table.rows[0].cells[j+1].text = f"{growth_rate:.1f}%"
                                sens_table.rows[0].cells[j+1].paragraphs[0].runs[0].font.size = Pt(8)
                                sens_table.rows[0].cells[j+1].paragraphs[0].runs[0].font.bold = True

                        # Data rows (discount rates)
                        for i, disc_delta in enumerate(disc_range):
                            if i >= len(results) or i+1 >= len(sens_table.rows):
                                continue

                            disc_rate = (sensitivity.get('baseDiscountRate', 0) + disc_delta) * 100
                            label = f"{disc_rate:.1f}%"
                            if sensitivity.get('method') == 'ndr':
                                label = f"Net {label}"
                            sens_table.rows[i+1].cells[0].text = label
                            sens_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                            sens_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True

                            result_row = results[i]
                            if isinstance(result_row, list):
                                for j, cell_data in enumerate(result_row):
                                    if j+1 < len(sens_table.rows[i+1].cells) and isinstance(cell_data, dict):
                                        total_pv = cell_data.get('totalPV', 0)
                                        sens_table.rows[i+1].cells[j+1].text = f"${total_pv:,.0f}"
                                        sens_table.rows[i+1].cells[j+1].paragraphs[0].runs[0].font.size = Pt(8)

                    doc.add_paragraph()

                    # Add detailed tables for each sensitivity scenario
                    if not tinari_mode:
                        doc.add_heading('TABLE 5: Detailed Sensitivity Scenarios (Year-by-Year)', level=2)
                        doc.add_paragraph("Complete year-by-year breakdown for each discount rate and growth rate combination showing full impact on damages").italic = True
                        doc.add_paragraph()

                        for i, disc_delta in enumerate(disc_range):
                            if i >= len(results):
                                continue

                            disc_rate = (sensitivity.get('baseDiscountRate', 0) + disc_delta) * 100
                            result_row = results[i]

                            if isinstance(result_row, list):
                                for j, cell_data in enumerate(result_row):
                                    if not isinstance(cell_data, dict):
                                        continue

                                    growth_delta = growth_range[j] if j < len(growth_range) else 0
                                    growth_rate = (sensitivity.get('baseGrowthRate', 0) + growth_delta) * 100
                                    scenario_schedule = cell_data.get('schedule', {})

                                    if not scenario_schedule or not scenario_schedule.get('rows'):
                                        continue

                                    # Add scenario heading
                                    scenario_title = f"Scenario: {disc_rate:.1f}% Discount, {growth_rate:.1f}% Growth"
                                    doc.add_heading(scenario_title, level=3)

                                    # Add totals summary
                                    totals_para = doc.add_paragraph()
                                    totals_para.add_run(f"Total PV: ${cell_data.get('totalPV', 0):,.2f}  |  ")
                                    totals_para.add_run(f"Past: ${cell_data.get('pastDam', 0):,.2f}  |  ")
                                    totals_para.add_run(f"Future PV: ${cell_data.get('futurePV', 0):,.2f}")
                                    totals_para.style = 'Normal'

                                    # Determine if UPS fringe
                                    use_ups_fringe_sens = cell_data.get('assumptions', {}).get('butFor', {}).get('fringeMethod') == 'ups'
                                    sens_options = cell_data.get('assumptions', {}).get('options', {})
                                    sens_aef = cell_data.get('assumptions', {}).get('aef', {})
                                    aef_on_sens = sens_options.get('includeAEF', False) and sens_aef.get('mode') == 'on'
                                    include_legals_sens = sens_options.get('includeLegals', True)
                                    show_fringe_sens = not aef_on_sens

                                    add_formula_note("Loss = BF After-Tax/AEF (+Fringe if shown) (+Legals if shown) - (ACT Earn + ACT Fringe (+ACT Legals if shown)); Past column shows past portion; PV(Future) and Survival-weighted values already computed for each cell.")

                                    hw_used_sens = pension_used_sens = fringe_total_used_sens = False
                                    fringe_used_sens = False
                                    if show_fringe_sens:
                                        if use_ups_fringe_sens:
                                            hw_used_sens = has_nonzero(scenario_schedule.get('rows', []), 'bfHW')
                                            pension_used_sens = has_nonzero(scenario_schedule.get('rows', []), 'bfPension')
                                            fringe_total_used_sens = has_nonzero(scenario_schedule.get('rows', []), 'bfFringe')
                                            fringe_used_sens = hw_used_sens or pension_used_sens or fringe_total_used_sens
                                        else:
                                            fringe_used_sens = has_nonzero(scenario_schedule.get('rows', []), 'bfFringe')
                                    legals_used_sens = include_legals_sens and has_nonzero(scenario_schedule.get('rows', []), 'bfLegals')
                                    act_fringe_used_sens = has_nonzero(scenario_schedule.get('rows', []), 'actFringe')
                                    act_legals_used_sens = include_legals_sens and has_nonzero(scenario_schedule.get('rows', []), 'actLegals')

                                    # Create detailed table
                                    sens_detail_headers = ['Year', 'Age', 'Portion', 'BF Gross', 'BF After-Tax / AEF']
                                    if fringe_used_sens:
                                        if use_ups_fringe_sens:
                                            if hw_used_sens:
                                                sens_detail_headers.append('H&W')
                                            if pension_used_sens:
                                                sens_detail_headers.append('Pension')
                                            if fringe_total_used_sens:
                                                sens_detail_headers.append('Total Fringe')
                                        else:
                                            sens_detail_headers.append('BF Fringe')
                                    if legals_used_sens:
                                        sens_detail_headers.append('BF Legally Req')
                                    sens_detail_headers.append('ACT Earn')
                                    if act_fringe_used_sens:
                                        sens_detail_headers.append('ACT Fringe')
                                    if act_legals_used_sens:
                                        sens_detail_headers.append('ACT Legals')
                                    sens_detail_headers.extend(['Loss', 'Past'])
                                    sens_detail_headers.extend(['Future (Raw)', 'Future (Survival)', 'PV(Future)', 'Survival Prob'])

                                    detail_table = doc.add_table(rows=1, cols=len(sens_detail_headers))
                                    detail_table.style = 'Light Grid Accent 1'

                                    # Header row
                                    hdr_cells = detail_table.rows[0].cells
                                    for idx, header in enumerate(sens_detail_headers):
                                        hdr_cells[idx].text = header
                                        if hdr_cells[idx].paragraphs and len(hdr_cells[idx].paragraphs) > 0:
                                            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
                                            hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(8)

                                    # Data rows
                                    for row_data in scenario_schedule.get('rows', []):
                                        if not isinstance(row_data, dict):
                                            continue

                                        row_cells = detail_table.add_row().cells
                                        values = [
                                            str(row_data.get('year', '')),
                                            str(row_data.get('age', '')),
                                            f"{row_data.get('portion', 0):.3f}" if isinstance(row_data.get('portion'), (int, float)) else str(row_data.get('portion', '')),
                                            f"${row_data.get('bfGross', 0):,.2f}",
                                            f"${row_data.get('bfAdj', 0):,.2f}"
                                        ]

                                        if fringe_used_sens:
                                            if use_ups_fringe_sens:
                                                if hw_used_sens:
                                                    values.append(f"${row_data.get('bfHW', 0):,.2f}")
                                                if pension_used_sens:
                                                    values.append(f"${row_data.get('bfPension', 0):,.2f}")
                                                if fringe_total_used_sens:
                                                    values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                                            else:
                                                values.append(f"${row_data.get('bfFringe', 0):,.2f}")

                                        if legals_used_sens:
                                            values.append(f"${row_data.get('bfLegals', 0):,.2f}")
                                        values.append(f"${row_data.get('actE', 0):,.2f}")
                                        if act_fringe_used_sens:
                                            values.append(f"${row_data.get('actFringe', 0):,.2f}")
                                        if act_legals_used_sens:
                                            values.append(f"${row_data.get('actLegals', 0):,.2f}")
                                        values.extend([
                                            f"${row_data.get('loss', 0):,.2f}",
                                            f"${row_data.get('pastPart', 0):,.2f}"
                                        ])

                                        values.extend([
                                            f"${row_data.get('futurePart', 0):,.2f}",
                                            f"${row_data.get('survivalWeightedFuture', row_data.get('futurePart', 0)):,.2f}",
                                            f"${row_data.get('pvFuture', 0):,.2f}",
                                            f"{(row_data.get('survivalProb') or 1):.3f}"
                                        ])

                                        for idx, value in enumerate(values):
                                            if idx < len(row_cells):
                                                row_cells[idx].text = value
                                                if row_cells[idx].paragraphs and len(row_cells[idx].paragraphs) > 0 and row_cells[idx].paragraphs[0].runs and len(row_cells[idx].paragraphs[0].runs) > 0:
                                                    row_cells[idx].paragraphs[0].runs[0].font.size = Pt(7)

                                    doc.add_paragraph()  # Spacing between scenarios
                    else:
                        doc.add_heading('Tinari Mode: Sensitivity Tables Omitted', level=3)
                        doc.add_paragraph('Full discount/growth coverage is captured in the matrix and heatmap above; per-scenario column grids are hidden for Tinari exports.').italic = True

                except Exception as e:
                    print(f"Error adding sensitivity analysis: {e}")
                    import traceback
                    traceback.print_exc()
                    # Continue with export even if this section fails
            elif not include_discounting:
                doc.add_heading('Sensitivity Analysis', level=2)
                doc.add_paragraph('Present value discounting is OFF, so sensitivity to discount rates is skipped.').italic = True

            # Add Summary Table (year-over-year)
            if schedule.get('rows'):
                try:
                    show_fringe = not aef_on
                    add_yoy_summary_table(
                        schedule.get('rows', []),
                        totals,
                        use_ups_fringe,
                        show_fringe,
                        include_legals,
                        'TABLE 6: Year-Over-Year Loss Summary',
                        'Condensed annual summary showing key earnings and loss figures for each year',
                        include_discounting
                    )
                except Exception as e:
                    print(f"Error adding summary table: {e}")
                    # Continue with export even if this section fails

            # ==================== ADDITIONAL COMPREHENSIVE TABLES ====================
            doc.add_heading('SUPPLEMENTARY ANALYSIS TABLES', level=1)
            doc.add_paragraph('Additional detailed breakdowns for comprehensive damages documentation')
            doc.add_paragraph()

            # Additional Table 1: Component Breakdown (But-For Components Only)
            if not tinari_mode and all_rows:
                try:
                    doc.add_heading('TABLE 7: But-For Earnings Components Breakdown', level=2)
                    doc.add_paragraph('Detailed breakdown of all but-for earning components by year').italic = True

                    add_formula_note("Total BF Package = BF After-Tax/AEF (+Fringe if shown) (+Legally Required if shown); Portion = fraction of the year included.")

                    comp_headers = ['Year', 'Age', 'BF Gross', 'BF After-Tax/AEF']
                    fringe_used = False
                    hw_used = pension_used = fringe_total_used = False
                    if show_fringe:
                        if use_ups_fringe:
                            hw_used = has_nonzero(all_rows, 'bfHW')
                            pension_used = has_nonzero(all_rows, 'bfPension')
                            fringe_total_used = has_nonzero(all_rows, 'bfFringe')
                            if hw_used:
                                comp_headers.append('H&W')
                                fringe_used = True
                            if pension_used:
                                comp_headers.append('Pension')
                                fringe_used = True
                            if fringe_total_used:
                                comp_headers.append('Total Fringe')
                                fringe_used = True
                        else:
                            fringe_used = has_nonzero(all_rows, 'bfFringe')
                            if fringe_used:
                                comp_headers.append('BF Fringe')
                    legals_used = include_legals and has_nonzero(all_rows, 'bfLegals')
                    if legals_used:
                        comp_headers.append('BF Legally Req')
                    comp_headers.extend(['Total BF Package', 'Portion'])

                    comp_table = doc.add_table(rows=1, cols=len(comp_headers))
                    comp_table.style = 'Light Grid Accent 1'

                    hdr_cells = comp_table.rows[0].cells
                    for i, header in enumerate(comp_headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

                    for row_data in all_rows:
                        if not isinstance(row_data, dict):
                            continue
                        row_cells = comp_table.add_row().cells

                        bf_total = row_data.get('bfAdj', 0)
                        values = [
                            str(row_data.get('year', '')),
                            str(row_data.get('age', '')),
                            f"${row_data.get('bfGross', 0):,.2f}",
                            f"${row_data.get('bfAdj', 0):,.2f}"
                        ]

                        if show_fringe:
                            if use_ups_fringe:
                                if hw_used:
                                    values.append(f"${row_data.get('bfHW', 0):,.2f}")
                                if pension_used:
                                    values.append(f"${row_data.get('bfPension', 0):,.2f}")
                                if fringe_total_used:
                                    values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                                bf_total += row_data.get('bfFringe', 0)
                            else:
                                if fringe_used:
                                    values.append(f"${row_data.get('bfFringe', 0):,.2f}")
                                bf_total += row_data.get('bfFringe', 0)

                        if legals_used:
                            values.append(f"${row_data.get('bfLegals', 0):,.2f}")
                            bf_total += row_data.get('bfLegals', 0)

                        values.extend([
                            f"${bf_total:,.2f}",
                            f"{row_data.get('portion', 0):.3f}" if isinstance(row_data.get('portion'), (int, float)) else str(row_data.get('portion', ''))
                        ])

                        for i, value in enumerate(values):
                            if i < len(row_cells):
                                row_cells[i].text = value
                                if row_cells[i].paragraphs and len(row_cells[i].paragraphs) > 0 and row_cells[i].paragraphs[0].runs and len(row_cells[i].paragraphs[0].runs) > 0:
                                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

                    doc.add_paragraph()
                except Exception as e:
                    print(f"Error adding component breakdown table: {e}")

            # Additional Table 2: Actual Earnings Components Breakdown
            if not tinari_mode and all_rows:
                try:
                    doc.add_heading('TABLE 8: Actual Earnings Components Breakdown', level=2)
                    doc.add_paragraph('Detailed breakdown of all actual/post-injury earning components by year').italic = True

                    add_formula_note("Total Actual Package = Actual Earnings + Actual Fringe (+ Legally Required if shown).")

                    act_headers = ['Year', 'Age', 'Actual Earnings']
                    act_fringe_used = has_nonzero(all_rows, 'actFringe')
                    act_legals_used = include_legals and has_nonzero(all_rows, 'actLegals')
                    if act_fringe_used:
                        act_headers.append('Actual Fringe')
                    if act_legals_used:
                        act_headers.append('Actual Legally Req')
                    act_headers.append('Total Actual Package')

                    act_table = doc.add_table(rows=1, cols=len(act_headers))
                    act_table.style = 'Light Grid Accent 1'

                    hdr_cells = act_table.rows[0].cells
                    for i, header in enumerate(act_headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

                    for row_data in all_rows:
                        if not isinstance(row_data, dict):
                            continue
                        row_cells = act_table.add_row().cells
                        act_total = row_data.get('actE', 0)
                        if act_fringe_used:
                            act_total += row_data.get('actFringe', 0)
                        if act_legals_used:
                            act_total += row_data.get('actLegals', 0)

                        values = [
                            str(row_data.get('year', '')),
                            str(row_data.get('age', '')),
                            f"${row_data.get('actE', 0):,.2f}"
                        ]
                        if act_fringe_used:
                            values.append(f"${row_data.get('actFringe', 0):,.2f}")
                        if act_legals_used:
                            values.append(f"${row_data.get('actLegals', 0):,.2f}")
                        values.append(f"${act_total:,.2f}")

                        for i, value in enumerate(values):
                            row_cells[i].text = value
                            if row_cells[i].paragraphs and len(row_cells[i].paragraphs) > 0 and row_cells[i].paragraphs[0].runs and len(row_cells[i].paragraphs[0].runs) > 0:
                                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

                    doc.add_paragraph()
                except Exception as e:
                    print(f"Error adding actual components table: {e}")

            # Additional Table 3: Present Value and Survival Analysis
            if not tinari_mode and all_rows and include_discounting:
                try:
                    doc.add_heading('TABLE 9: Present Value and Survival Probability Analysis', level=2)
                    doc.add_paragraph('Shows the impact of present value discounting and survival probabilities on future damages').italic = True

                    add_formula_note("Future (Survival) = Future (Raw) × Survival Prob; PV(Future) = Future (Survival) discounted to valuation date.")
    
                    pv_table = doc.add_table(rows=1, cols=7)
                    pv_table.style = 'Light Grid Accent 1'
    
                    pv_headers = ['Year', 'Age', 'Annual Loss', 'Survival Prob', 'Future (Raw)', 'Future (Survival)', 'PV(Future)']
                    hdr_cells = pv_table.rows[0].cells
                    for i, header in enumerate(pv_headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
                    for row_data in all_rows:
                        if not isinstance(row_data, dict):
                            continue
                        row_cells = pv_table.add_row().cells
    
                        values = [
                            str(row_data.get('year', '')),
                            str(row_data.get('age', '')),
                            f"${row_data.get('loss', 0):,.2f}",
                            f"{(row_data.get('survivalProb') or 1):.4f}",
                            f"${row_data.get('futurePart', 0):,.2f}",
                            f"${row_data.get('survivalWeightedFuture', row_data.get('futurePart', 0)):,.2f}",
                            f"${row_data.get('pvFuture', 0):,.2f}"
                        ]
    
                        for i, value in enumerate(values):
                            row_cells[i].text = value
                            if row_cells[i].paragraphs and len(row_cells[i].paragraphs) > 0 and row_cells[i].paragraphs[0].runs and len(row_cells[i].paragraphs[0].runs) > 0:
                                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    
                    doc.add_paragraph()
                except Exception as e:
                    print(f"Error adding PV/survival table: {e}")
            elif not tinari_mode and all_rows and not include_discounting:
                doc.add_heading('TABLE 9: Present Value and Survival Probability Analysis', level=2)
                doc.add_paragraph('Present value discounting is OFF, so this table is omitted.').italic = True

            # Additional Table 5: Loss Components Comparison
            if not tinari_mode and all_rows:
                try:
                    doc.add_heading('TABLE 11: Annual Loss Components Comparison', level=2)
                    doc.add_paragraph('Compares but-for total compensation to actual total compensation to show annual loss').italic = True
    
                    add_formula_note("Annual Loss = But-For Total Compensation - Actual Total Compensation; Loss % = Annual Loss / But-For Total Compensation.")
    
                    loss_table = doc.add_table(rows=1, cols=6)
                    loss_table.style = 'Light Grid Accent 1'
    
                    loss_headers = ['Year', 'Age', 'But-For Total', 'Actual Total', 'Annual Loss', 'Loss %']
                    hdr_cells = loss_table.rows[0].cells
                    for i, header in enumerate(loss_headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
                    for row_data in all_rows:
                        if not isinstance(row_data, dict):
                            continue
                        row_cells = loss_table.add_row().cells
    
                        bf_total = (row_data.get('bfAdj', 0) + row_data.get('bfFringe', 0) +
                                   row_data.get('bfLegals', 0))
                        act_total = (row_data.get('actE', 0) + row_data.get('actFringe', 0) +
                                    row_data.get('actLegals', 0))
                        loss = row_data.get('loss', 0)
                        loss_pct = (loss / bf_total * 100) if bf_total > 0 else 0
    
                        values = [
                            str(row_data.get('year', '')),
                            str(row_data.get('age', '')),
                            f"${bf_total:,.2f}",
                            f"${act_total:,.2f}",
                            f"${loss:,.2f}",
                            f"{loss_pct:.1f}%"
                        ]
    
                        for i, value in enumerate(values):
                            row_cells[i].text = value
                            if row_cells[i].paragraphs and len(row_cells[i].paragraphs) > 0 and row_cells[i].paragraphs[0].runs and len(row_cells[i].paragraphs[0].runs) > 0:
                                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    
                    doc.add_paragraph()
                except Exception as e:
                    print(f"Error adding loss components table: {e}")

            # Move all visuals to the end for easier reading
            doc.add_page_break()
            doc.add_heading('Executive Visual Summary', level=1)
            doc.add_paragraph('Charts and graphs appear after all tables for a simpler, story-first read.').italic = True
            doc.add_paragraph()

            # Chart 1: Damages Breakdown Pie Chart
            try:
                pie_chart = create_damages_breakdown_pie(totals, "Total Economic Damages Breakdown")
                if pie_chart:
                    doc.add_heading('Total Damages Summary', level=2)
                    doc.add_picture(pie_chart, width=Inches(6))
                    doc.add_paragraph()
            except Exception as e:
                print(f"Error creating damages pie chart: {e}")

            # Chart 2: Annual Loss Bar Chart
            try:
                if all_rows:
                    loss_chart = create_annual_loss_chart(all_rows, "Annual Economic Loss by Year")
                    if loss_chart:
                        doc.add_heading('Annual Economic Losses', level=2)
                        doc.add_paragraph('This chart shows the economic loss for each year. Red bars indicate losses (but-for scenario exceeds actual earnings).')
                        doc.add_picture(loss_chart, width=Inches(7))
                        doc.add_paragraph()
            except Exception as e:
                print(f"Error creating annual loss chart: {e}")

            # Chart 3: But-For vs Actual Earnings Comparison
            try:
                if all_rows:
                    earnings_chart = create_earnings_comparison_chart(all_rows, "But-For vs Actual Earnings Comparison")
                    if earnings_chart:
                        doc.add_heading('Earnings Trajectory Comparison', level=2)
                        doc.add_paragraph('Green line shows projected but-for earnings. Red line shows actual/post-injury earnings. Shaded area represents the economic loss.')
                        doc.add_picture(earnings_chart, width=Inches(7))
                        doc.add_paragraph()
            except Exception as e:
                print(f"Error creating earnings comparison chart: {e}")

            # Chart 4: Cumulative Damages
            try:
                if all_rows:
                    cumulative_chart = create_cumulative_damages_chart(all_rows, "Cumulative Economic Damages Over Time")
                    if cumulative_chart:
                        doc.add_heading('Cumulative Damages Over Time', level=2)
                        doc.add_paragraph('This area chart shows how damages accumulate over time. Pink area represents past damages; blue area represents future damages (present value).')
                        doc.add_picture(cumulative_chart, width=Inches(7))
                        doc.add_paragraph()
            except Exception as e:
                print(f"Error creating cumulative damages chart: {e}")

            # Chart 5: Sensitivity Heatmap (visual only)
            try:
                if include_discounting and sensitivity and sensitivity.get('results'):
                    sens_heatmap = create_sensitivity_heatmap(sensitivity)
                    if sens_heatmap:
                        doc.add_heading('Sensitivity Analysis Visual Heatmap', level=2)
                        doc.add_paragraph('Total present value across the discount/growth grid; brighter green = higher damages, red = lower.')
                        doc.add_picture(sens_heatmap, width=Inches(7))
                        doc.add_paragraph()
            except Exception as e:
                print(f"Error creating sensitivity heatmap: {e}")

            # Chart 6: UPS Fringe Breakdown (if applicable)
            try:
                if all_rows and use_ups_fringe:
                    ups_chart = create_ups_fringe_breakdown_chart(all_rows, "UPS Fringe Benefits: Health & Welfare vs Pension")
                    if ups_chart:
                        doc.add_heading('UPS-Specific Fringe Benefits Breakdown', level=2)
                        doc.add_paragraph('Detailed side-by-side comparison of UPS Health & Welfare contributions and Pension contributions by year.')
                        doc.add_picture(ups_chart, width=Inches(7))
                        doc.add_paragraph()
            except Exception as e:
                print(f"Error creating UPS fringe breakdown chart: {e}")

            # Jury-friendly summary charts (bottom of report, one simple graphic each)
            try:
                doc.add_heading('Plain-English Jury Visuals', level=1)
                doc.add_paragraph('Simple one-frame charts: what was lost, how long it lasts, the growth/inflation factor, and total loss.').italic = True
    
                jury_items = create_jury_items_chart(totals, "What Was Lost")
                if jury_items:
                    doc.add_heading('What Was Lost', level=2)
                    doc.add_picture(jury_items, width=Inches(6))
                    doc.add_paragraph()
    
                jury_years = create_jury_years_chart(schedule, "How Long The Loss Lasts")
                if jury_years:
                    doc.add_heading('How Long The Loss Lasts', level=2)
                    doc.add_picture(jury_years, width=Inches(5))
                    doc.add_paragraph()
    
                jury_growth = create_jury_growth_chart(assumptions, "Growth / Inflation Factor Used")
                if jury_growth:
                    doc.add_heading('Growth / Inflation Factor Used', level=2)
                    doc.add_picture(jury_growth, width=Inches(6))
                    doc.add_paragraph()
    
                jury_total = create_jury_total_chart(totals, "Total Economic Loss (Present Value)")
                if jury_total:
                    doc.add_heading('Total Economic Loss (PV)', level=2)
                    doc.add_picture(jury_total, width=Inches(5))
                    doc.add_paragraph()
            except Exception as e:
                print(f"Error creating jury breakdown charts: {e}")
    
            # Save to BytesIO
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
    
            # Generate filename
            filename = f"damages_report_{case_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/export/excel', methods=['POST'])
    def export_excel():
        """Generate Excel workbook with Tinari-style tables and formulas"""
        try:
            data = request.json
            assumptions = data.get('assumptions', {})
            schedule = data.get('schedule', {})
            retirement_scenarios = data.get('retirementScenarios', [])
            sensitivity = data.get('sensitivityAnalysis', {})
            validation_notes = _validate_assumption_ranges(assumptions)
            provenance = _build_provenance_metadata(assumptions)

            wb = openpyxl.Workbook()
            base_sheet = wb.active
            base_sheet.title = 'Tinari Base'

            def fmt_pct(val):
                return f"{(val or 0)*100:.2f}%" if isinstance(val, (int, float)) else ''

            def add_tinari_sheet(name, sched, assumps, ws=None):
                if ws is None:
                    safe_title = (name or 'Sheet')
                    if not isinstance(safe_title, str):
                        safe_title = 'Sheet'
                    safe_title = safe_title[:31] or 'Sheet'
                    ws = wb.create_sheet(safe_title)
                ws.append(['AEF (Adjusted Earnings Factor)', assumps.get('aef', {}).get('factor', '')])
                ws.append(['GE (Gross Earnings Base)', assumps.get('aef', {}).get('grossEarningsBase', '')])
                ws.append(['WLE (Worklife Adjusted Earnings Base)', assumps.get('aef', {}).get('wle', '')])
                ws.append(['UF (Unemployment Factor)', assumps.get('aef', {}).get('ufEff', '')])
                ws.append(['TR (Combined Effective Tax Rate)', assumps.get('aef', {}).get('tlEff', '')])
                ws.append(['FB (Fringe Benefits Loading)', assumps.get('aef', {}).get('fringePct', '')])

                growth_label = ''
                if assumps.get('butFor', {}).get('growthMethod') == 'fixed':
                    growth_label = fmt_pct(assumps.get('butFor', {}).get('growth', 0)) + ' growth'
                elif assumps.get('butFor', {}).get('growthMethod') == 'ups':
                    growth_label = 'UPS contract growth (varies)'
                elif assumps.get('butFor', {}).get('growthMethod') == 'series':
                    growth_label = 'Series growth (varies)'

                discount_rate = (assumps.get('discount', {}).get('ndr') if assumps.get('discount', {}).get('method') == 'ndr' else assumps.get('discount', {}).get('rate', 0)) or 0
                ws.append(['Discount Rate', discount_rate, growth_label])
                discount_cell = ws.cell(ws.max_row, 2)
                discount_abs = f"${get_column_letter(discount_cell.column)}{discount_cell.row}"

                ws.append([])  # spacer
                headers = ['Year', 'Age', 'Portion of Year', 'Base Earnings', 'Adjusted Income', 'Present Value', 'Years From Valuation']
                ws.append(headers)

                # Algebraic audit helpers
                ws.append(['', '', '', 'Adjusted = Base * AEF', 'PV = Adjusted / (1 + r)^years_from_val', '', ''])
                for cell in ws[ws.max_row]:
                    cell.font = Font(italic=True)

                first_data_row = None
                past_start = past_end = None
                future_start = future_end = None
                past_total_row = future_total_row = None

                val_dt = None
                try:
                    val_date = assumps.get('dates', {}).get('valuation')
                    val_dt = datetime.fromisoformat(val_date) if val_date else None
                except Exception:
                    val_dt = None

                def years_from_val(year):
                    if year is None:
                        return 0
                    try:
                        y_int = int(year)
                    except Exception:
                        return 0
                    if not val_dt:
                        return 0
                    mid = datetime(y_int, 7, 1)
                    return max(0, (mid - val_dt).days / 365.25)

                def add_section(label):
                    # Add a full-width row so we don't access missing cells
                    ws.append([label] + [''] * (len(headers) - 1))
                    row = ws[ws.max_row]
                    if row:
                        row[0].font = Font(bold=True)

                rows_pre = sched.get('rowsPre', []) or []
                rows_post = sched.get('rowsPost', []) or []
                val_year = val_dt.year if val_dt else None

                add_section('Past Years')
                for r in rows_pre:
                    yrs = years_from_val(r.get('year'))
                    adjusted_row = ws.max_row + 1
                    if first_data_row is None:
                        first_data_row = adjusted_row
                    if past_start is None:
                        past_start = adjusted_row
                    y = r.get('year')
                    pv_combined = (r.get('pastPart', 0) or 0) + (r.get('pvFuture', 0) or 0)
                    ws.append([
                        r.get('year', ''),
                        r.get('age', ''),
                        r.get('portion', 0),
                        r.get('bfGross', 0),
                        f"=D{adjusted_row}*$B$1" if ws.cell(1,2).value else r.get('bfAdj', 0),
                        pv_combined,
                        yrs
                    ])

                past_end = ws.max_row if past_start else None
                if past_start and past_end and past_end >= past_start:
                    past_total_row = ws.max_row + 1
                    ws.append(['', '', '', 'Past Totals', f"=SUM(E{past_start}:E{past_end})", f"=SUM(F{past_start}:F{past_end})", ''])
                    for col_idx in range(4, 7):
                        ws.cell(past_total_row, col_idx).font = Font(bold=True)

                add_section('Future Years')
                for r in rows_post:
                    yrs = years_from_val(r.get('year'))
                    adjusted_row = ws.max_row + 1
                    if first_data_row is None:
                        first_data_row = adjusted_row
                    if future_start is None:
                        future_start = adjusted_row
                    y = r.get('year')
                    pv_combined = (r.get('pastPart', 0) or 0) + (r.get('pvFuture', 0) or 0)
                    ws.append([
                        r.get('year', ''),
                        r.get('age', ''),
                        r.get('portion', 0),
                        r.get('bfGross', 0),
                        f"=D{adjusted_row}*$B$1" if ws.cell(1,2).value else r.get('bfAdj', 0),
                        pv_combined,
                        yrs
                    ])

                future_end = ws.max_row if future_start else None
                if future_start and future_end and future_end >= future_start:
                    future_total_row = ws.max_row + 1
                    ws.append(['', '', '', 'Future Totals', f"=SUM(E{future_start}:E{future_end})", f"=SUM(F{future_start}:F{future_end})", ''])
                    for col_idx in range(4, 7):
                        ws.cell(future_total_row, col_idx).font = Font(bold=True)

                if first_data_row:
                    data_last_row = future_end or past_end
                    if data_last_row is None:
                        data_last_row = ws.max_row

                    total_row = ws.max_row + 1

                    fv_refs = []
                    pv_refs = []
                    if past_total_row:
                        fv_refs.append(f"E{past_total_row}")
                        pv_refs.append(f"F{past_total_row}")
                    if future_total_row:
                        fv_refs.append(f"E{future_total_row}")
                        pv_refs.append(f"F{future_total_row}")

                    total_future_formula = f"=SUM({','.join(fv_refs)})" if fv_refs else f"=SUM(E{first_data_row}:E{data_last_row})"
                    total_pv_formula = f"=SUM({','.join(pv_refs)})" if pv_refs else f"=SUM(F{first_data_row}:F{data_last_row})"

                    ws.append(['', '', '', 'Total', total_future_formula, total_pv_formula, ''])
                    for col_idx in range(4, 7):
                        ws.cell(total_row, col_idx).font = Font(bold=True)

                ws.column_dimensions[get_column_letter(7)].hidden = True
                for col in range(1, 8):
                    ws.column_dimensions[get_column_letter(col)].width = 18

            add_tinari_sheet('Tinari Base', schedule, assumptions, base_sheet)

            if isinstance(retirement_scenarios, list):
                for idx, scenario in enumerate(retirement_scenarios):
                    sched = scenario.get('schedule', {})
                    assumps = scenario.get('assumptions', assumptions)
                    name = scenario.get('name', f'Retirement {idx+1}')
                    add_tinari_sheet(f"Retire {idx+1}", sched, assumps)

            if sensitivity and sensitivity.get('results'):
                ws = wb.create_sheet('Sensitivity')
                ws.append(['Discount Δ', 'Growth Δ', 'Total PV', 'Past', 'Future PV'])
                disc_range = sensitivity.get('discountRange', sensitivity.get('range', []))
                growth_range = sensitivity.get('growthRange', sensitivity.get('range', []))
                results = sensitivity.get('results', [])
                for i, disc in enumerate(disc_range):
                    row_results = results[i] if i < len(results) else []
                    for j, cell in enumerate(row_results):
                        if not isinstance(cell, dict):
                            continue
                        growth_delta = growth_range[j] if j < len(growth_range) else 0
                        ws.append([disc, growth_delta, cell.get('totalPV', 0), cell.get('pastDam', 0), cell.get('futurePV', 0)])

            prov_sheet = wb.create_sheet('Provenance')
            prov_sheet.append(['Generated (UTC)', provenance['generated_at']])
            prov_sheet.append(['Assumptions fingerprint (SHA-256)', provenance['fingerprint']])
            prov_sheet.append([])
            prov_sheet.append(['Sources'])
            if provenance['sources']:
                for source in provenance['sources']:
                    prov_sheet.append(['', source])
            else:
                prov_sheet.append(['', 'No explicit sources provided.'])
            prov_sheet.append([])
            prov_sheet.append(['Validation'])
            if validation_notes:
                for warning in validation_notes:
                    prov_sheet.append(['', warning])
            else:
                prov_sheet.append(['', 'All monitored inputs fall within configured ranges.'])
            prov_sheet.column_dimensions['A'].width = 36
            prov_sheet.column_dimensions['B'].width = 120

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
            wb.save(tmp.name)
            tmp.seek(0)
            return send_file(tmp.name, as_attachment=True, download_name='damages_report.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/stats', methods=['GET'])
    def get_stats():
        """Get database statistics"""
        return jsonify({
            'success': True,
            'stats': {
                'evaluees': Evaluee.query.count(),
                'cases': Case.query.count(),
                'calculations': Calculation.query.count()
            }
        })

    # ==================== STATIC FILES ====================
    # Serve the SPA from the repo-level /static directory, regardless of cwd.
    static_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static'))

    @app.route('/', defaults={'path': ''})
    @app.route('/<path:path>')
    def static_files(path):
        """Serve compiled SPA assets with index.html fallback"""
        target_path = os.path.join(static_dir, path)
        if path and os.path.exists(target_path):
            return send_from_directory(static_dir, path)
        # Fallback to SPA shell
        return send_from_directory(static_dir, 'index.html')

    # ==================== ERROR HANDLERS ====================

    @app.errorhandler(404)
    def not_found(error):
        return jsonify({
            'success': False,
            'error': 'Resource not found'
        }), 404

    @app.errorhandler(500)
    def internal_error(error):
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': 'Internal server error'
        }), 500

    return app


if __name__ == '__main__':
    app = create_app(os.environ.get('FLASK_ENV', 'development'))
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
